"""
AI 聊天智能体 - 后端服务
支持创建自定义智能体并进行对话
"""
import json
import os
import re
import uuid
import threading
import time
import hashlib
import secrets
import sqlite3
import urllib.request
import urllib.parse
import io
import shutil
import logging
from datetime import datetime
from pathlib import Path
from typing import Optional, List

from fastapi import FastAPI, HTTPException, UploadFile, File, Header, Form, Depends, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
from openai import OpenAI
from dotenv import load_dotenv

# ==================== 配置 ====================
BASE_DIR = Path(__file__).parent

# 加载 .env 文件
load_dotenv(BASE_DIR / ".env")

DB_PATH = BASE_DIR / "data" / "agents.db"
SYSTEM_USER_ID = "00000000-0000-0000-0000-000000000000"  # 系统内置智能体归属
DEFAULT_AVATARS = ["🤖", "🧠", "🎭", "🌟", "🦊", "🐱", "🐶", "🦄", "🐲", "👾"]

# 从环境变量读取 API 配置
API_KEY = os.getenv("OPENAI_API_KEY", "")
API_BASE = os.getenv("OPENAI_API_BASE", "https://api.siliconflow.cn/v1")
MODEL_NAME = os.getenv("OPENAI_MODEL", "deepseek-ai/DeepSeek-V3")
OPENWEATHER_API_KEY = os.getenv("OPENWEATHER_API_KEY", "")
ADMIN_USERNAME = os.getenv("ADMIN_USERNAME", "").strip()  # 管理员用户名

# ==================== RAG / 知识库配置 ====================
TOP_K_RESULTS = int(os.getenv("TOP_K_RESULTS", "5"))
CHUNK_SIZE = int(os.getenv("CHUNK_SIZE", "1000"))
CHUNK_OVERLAP = int(os.getenv("CHUNK_OVERLAP", "200"))
EMBEDDING_MODEL = os.getenv("EMBEDDING_MODEL", "BAAI/bge-large-zh-v1.5")
CHROMA_DIR = BASE_DIR / "data" / "chroma"

# 支持的知识库文件类型
ALLOWED_DOC_TYPES = {
    "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "txt": "text/plain",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "xls": "application/vnd.ms-excel",
}

# 日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# ==================== 数据库 ====================
os.makedirs(BASE_DIR / "data", exist_ok=True)


def get_db():
    """获取数据库连接（每次创建新连接，线程安全）"""
    conn = sqlite3.connect(str(DB_PATH))
    conn.row_factory = sqlite3.Row
    return conn


def _migrate_tables_for_user_id(conn):
    """迁移旧的共享数据库，为所有表添加 user_id 字段"""
    tables = ["agents", "conversations", "knowledge_files", "reminders", "user_memories"]
    for tbl in tables:
        try:
            conn.execute(f"ALTER TABLE {tbl} ADD COLUMN user_id TEXT NOT NULL DEFAULT ''")
            conn.commit()
        except sqlite3.OperationalError:
            pass  # 字段已存在

    # 为旧数据填充：如果 user_memories 的 UNIQUE 约束不含 user_id，需要重建
    try:
        existing = conn.execute("SELECT sql FROM sqlite_master WHERE type='table' AND name='user_memories'").fetchone()
        if existing and "UNIQUE(agent_id, category, key)" in existing["sql"]:
            conn.execute("""
                CREATE TABLE user_memories_new2 (
                    id TEXT PRIMARY KEY,
                    user_id TEXT NOT NULL DEFAULT '',
                    agent_id TEXT NOT NULL DEFAULT '',
                    category TEXT NOT NULL DEFAULT 'general',
                    key TEXT NOT NULL,
                    value TEXT NOT NULL,
                    source_conv_id TEXT,
                    confidence REAL NOT NULL DEFAULT 0.5,
                    created_at TEXT NOT NULL DEFAULT (datetime('now')),
                    updated_at TEXT NOT NULL DEFAULT (datetime('now')),
                    UNIQUE(user_id, agent_id, category, key)
                )
            """)
            conn.execute("INSERT INTO user_memories_new2 SELECT id, '', agent_id, category, key, value, source_conv_id, confidence, created_at, updated_at FROM user_memories")
            conn.execute("DROP TABLE user_memories")
            conn.execute("ALTER TABLE user_memories_new2 RENAME TO user_memories")
            conn.commit()
    except Exception:
        pass  # 迁移已完成或不需要

    # 🔥 多用户隔离：将系统内置智能体的 user_id 设为 SYSTEM_USER_ID
    try:
        conn.execute(
            "UPDATE agents SET user_id = ? WHERE user_id = ''",
            (SYSTEM_USER_ID,)
        )
        conn.commit()
        logger.info(f"已将旧智能体的 user_id 更新为 {SYSTEM_USER_ID}（系统共享）")
    except Exception:
        pass

    # 🔥 清理旧的无主对话数据（user_id='' 的旧数据无法归属到任何用户，安全清理）
    try:
        deleted = conn.execute("DELETE FROM reminders WHERE user_id = ''").rowcount
        deleted += conn.execute("DELETE FROM knowledge_files WHERE user_id = ''").rowcount
        deleted += conn.execute("DELETE FROM user_memories WHERE user_id = ''").rowcount
        deleted += conn.execute("DELETE FROM messages WHERE conversation_id IN (SELECT id FROM conversations WHERE user_id = '')").rowcount
        deleted += conn.execute("DELETE FROM conversations WHERE user_id = ''").rowcount
        conn.commit()
        logger.info(f"已清理 {deleted} 条无主旧数据，多用户隔离已生效")
    except Exception:
        pass


def init_db():
    """初始化数据库表"""
    conn = get_db()
    conn.executescript("""
        CREATE TABLE IF NOT EXISTS agents (
            id TEXT PRIMARY KEY,
            user_id TEXT NOT NULL DEFAULT '',
            name TEXT NOT NULL,
            personality TEXT NOT NULL DEFAULT '',
            speaking_style TEXT NOT NULL DEFAULT '',
            character_setting TEXT NOT NULL DEFAULT '',
            avatar TEXT NOT NULL DEFAULT '🤖',
            system_prompt TEXT NOT NULL DEFAULT '',
            enable_search INTEGER NOT NULL DEFAULT 0,
            created_at TEXT NOT NULL DEFAULT (datetime('now')),
            updated_at TEXT NOT NULL DEFAULT (datetime('now'))
        );

        CREATE TABLE IF NOT EXISTS conversations (
            id TEXT PRIMARY KEY,
            agent_id TEXT NOT NULL,
            user_id TEXT NOT NULL DEFAULT '',
            title TEXT NOT NULL DEFAULT '新对话',
            created_at TEXT NOT NULL DEFAULT (datetime('now')),
            updated_at TEXT NOT NULL DEFAULT (datetime('now')),
            FOREIGN KEY (agent_id) REFERENCES agents(id) ON DELETE CASCADE
        );

        CREATE TABLE IF NOT EXISTS messages (
            id TEXT PRIMARY KEY,
            conversation_id TEXT NOT NULL,
            role TEXT NOT NULL CHECK(role IN ('user', 'assistant', 'system')),
            content TEXT NOT NULL,
            created_at TEXT NOT NULL DEFAULT (datetime('now')),
            FOREIGN KEY (conversation_id) REFERENCES conversations(id) ON DELETE CASCADE
        );

        -- 用户记忆表：存储从对话中提取的用户偏好/信息（按智能体隔离）
        CREATE TABLE IF NOT EXISTS user_memories (
            id TEXT PRIMARY KEY,
            user_id TEXT NOT NULL DEFAULT '',
            agent_id TEXT NOT NULL DEFAULT '',
            category TEXT NOT NULL DEFAULT 'general',
            key TEXT NOT NULL,
            value TEXT NOT NULL,
            source_conv_id TEXT,
            confidence REAL NOT NULL DEFAULT 0.5,
            created_at TEXT NOT NULL DEFAULT (datetime('now')),
            updated_at TEXT NOT NULL DEFAULT (datetime('now')),
            UNIQUE(user_id, agent_id, category, key)
        );

        -- 偏好提取日志：记录每次提取操作
        CREATE TABLE IF NOT EXISTS extraction_log (
            id TEXT PRIMARY KEY,
            user_id TEXT NOT NULL DEFAULT '',
            conversation_id TEXT NOT NULL,
            extracted_count INTEGER NOT NULL DEFAULT 0,
            extracted_summary TEXT NOT NULL DEFAULT '',
            created_at TEXT NOT NULL DEFAULT (datetime('now'))
        );

        -- 用户表：账号系统
        CREATE TABLE IF NOT EXISTS users (
            id TEXT PRIMARY KEY,
            username TEXT UNIQUE NOT NULL,
            password_hash TEXT NOT NULL,
            salt TEXT NOT NULL,
            nickname TEXT NOT NULL DEFAULT '',
            avatar_url TEXT NOT NULL DEFAULT '',
            token TEXT NOT NULL DEFAULT '',
            is_admin INTEGER NOT NULL DEFAULT 0,
            created_at TEXT NOT NULL DEFAULT (datetime('now')),
            updated_at TEXT NOT NULL DEFAULT (datetime('now'))
        );

        -- 知识库文件表
        CREATE TABLE IF NOT EXISTS knowledge_files (
            id TEXT PRIMARY KEY,
            agent_id TEXT NOT NULL,
            user_id TEXT NOT NULL DEFAULT '',
            filename TEXT NOT NULL,
            original_name TEXT NOT NULL,
            file_type TEXT NOT NULL DEFAULT 'txt',
            chunk_count INTEGER NOT NULL DEFAULT 0,
            file_size INTEGER NOT NULL DEFAULT 0,
            created_at TEXT NOT NULL DEFAULT (datetime('now')),
            FOREIGN KEY (agent_id) REFERENCES agents(id) ON DELETE CASCADE
        );

        -- 提醒任务表：当用户要求智能体定时提醒时使用
        CREATE TABLE IF NOT EXISTS reminders (
            id TEXT PRIMARY KEY,
            agent_id TEXT NOT NULL,
            user_id TEXT NOT NULL DEFAULT '',
            conversation_id TEXT NOT NULL,
            remind_at TEXT NOT NULL,
            content TEXT NOT NULL,
            status TEXT NOT NULL DEFAULT 'pending',
            sent_at TEXT,
            created_at TEXT NOT NULL DEFAULT (datetime('now')),
            FOREIGN KEY (agent_id) REFERENCES agents(id) ON DELETE CASCADE
        );
    """)
    conn.commit()

    # 迁移：为旧数据库补齐 character_setting 字段
    try:
        conn.execute("ALTER TABLE agents ADD COLUMN character_setting TEXT NOT NULL DEFAULT ''")
        conn.commit()
    except sqlite3.OperationalError:
        pass  # 字段已存在，跳过

    # 迁移：为旧数据库补齐 enable_search 字段
    try:
        conn.execute("ALTER TABLE agents ADD COLUMN enable_search INTEGER NOT NULL DEFAULT 0")
        conn.commit()
    except sqlite3.OperationalError:
        pass  # 字段已存在，跳过

    # 迁移：为所有表添加 user_id 字段（多用户隔离）
    _migrate_tables_for_user_id(conn)

    # 迁移：为 user_memories 表添加 agent_id 字段并更新唯一约束
    try:
        conn.execute("ALTER TABLE user_memories ADD COLUMN agent_id TEXT NOT NULL DEFAULT ''")
        conn.commit()
    except sqlite3.OperationalError:
        pass  # 字段已存在

    # 迁移：重建 user_memories 的唯一约束（包含 agent_id）
    try:
        # 检查旧约束是否还存在
        existing = conn.execute("SELECT sql FROM sqlite_master WHERE type='table' AND name='user_memories'").fetchone()
        if existing and "UNIQUE(category, key)" in existing["sql"]:
            conn.execute("""
                CREATE TABLE user_memories_new (
                    id TEXT PRIMARY KEY,
                    agent_id TEXT NOT NULL DEFAULT '',
                    category TEXT NOT NULL DEFAULT 'general',
                    key TEXT NOT NULL,
                    value TEXT NOT NULL,
                    source_conv_id TEXT,
                    confidence REAL NOT NULL DEFAULT 0.5,
                    created_at TEXT NOT NULL DEFAULT (datetime('now')),
                    updated_at TEXT NOT NULL DEFAULT (datetime('now')),
                    UNIQUE(agent_id, category, key)
                )
            """)
            conn.execute("INSERT INTO user_memories_new SELECT id, agent_id, category, key, value, source_conv_id, confidence, created_at, updated_at FROM user_memories")
            conn.execute("DROP TABLE user_memories")
            conn.execute("ALTER TABLE user_memories_new RENAME TO user_memories")
            conn.commit()
    except Exception:
        pass  # 迁移已执行或不需要

    # 迁移：添加 users 表（老数据库可能没有）
    try:
        conn.execute("""
            CREATE TABLE IF NOT EXISTS users (
                id TEXT PRIMARY KEY,
                username TEXT UNIQUE NOT NULL,
                password_hash TEXT NOT NULL,
                salt TEXT NOT NULL,
                nickname TEXT NOT NULL DEFAULT '',
                avatar_url TEXT NOT NULL DEFAULT '',
                token TEXT NOT NULL DEFAULT '',
                is_admin INTEGER NOT NULL DEFAULT 0,
                created_at TEXT NOT NULL DEFAULT (datetime('now')),
                updated_at TEXT NOT NULL DEFAULT (datetime('now'))
            )
        """)
        conn.commit()
    except Exception:
        pass

    # 迁移：添加 is_admin 列（老数据库可能没有此列）
    try:
        conn.execute("ALTER TABLE users ADD COLUMN is_admin INTEGER NOT NULL DEFAULT 0")
        conn.commit()
    except Exception:
        pass  # 列已存在或不可 ALTER，忽略

    # 迁移：添加 knowledge_files 表（知识库文件）
    try:
        conn.execute("""
            CREATE TABLE IF NOT EXISTS knowledge_files (
                id TEXT PRIMARY KEY,
                agent_id TEXT NOT NULL,
                filename TEXT NOT NULL,
                original_name TEXT NOT NULL,
                file_type TEXT NOT NULL DEFAULT 'txt',
                chunk_count INTEGER NOT NULL DEFAULT 0,
                file_size INTEGER NOT NULL DEFAULT 0,
                created_at TEXT NOT NULL DEFAULT (datetime('now')),
                FOREIGN KEY (agent_id) REFERENCES agents(id) ON DELETE CASCADE
            )
        """)
        conn.commit()
    except Exception:
        pass

    # 迁移：extraction_log 添加 user_id 列（老数据库可能没有此列）
    try:
        conn.execute("ALTER TABLE extraction_log ADD COLUMN user_id TEXT NOT NULL DEFAULT ''")
        conn.commit()
    except Exception:
        pass  # 列已存在

    # 检查是否有默认智能体
    cursor = conn.execute("SELECT COUNT(*) FROM agents")
    if cursor.fetchone()[0] == 0:
        _create_default_agents(conn)

    # 保障私人管家系列始终存在（即使已有其他智能体）
    _ensure_default_butlers(conn)

    conn.close()


def _create_default_agents(conn):
    """创建默认智能体（仅两位管家，其余由用户自由创建）"""
    defaults = [
        {
            "id": "butler-gentle-001",
            "name": "温和管家",
            "personality": "温和、细心、有条理、善于规划和共情",
            "speaking_style": "说话温和体贴、条理清晰，语气像一位值得信赖的私人管家。会用'我帮您看看'、'我为您整理一下'、'需要我帮忙吗'等句式，偶尔会加上'~'让语气更柔和。在给出建议时总是先肯定用户，再提出方案。",
            "character_setting": "你是本平台的私人管家，负责帮助用户管理文件和日程。\n\n你的定位：你是'私人管家服务'的温和风格代表。你的职责是协助用户把各种文档（PDF、Word、Excel、TXT等）上传到任意智能体的知识库中，帮用户规划日程、提炼文档要点、整理信息。用户可以和本平台上任何智能体对话，给任何智能体上传文件作为知识库——而你是最温柔贴心的那一位。\n\n本平台特色：\n- 用户可以把工作文档、学习资料、合同等上传给任意智能体，该智能体会自动检索文档内容来回答用户问题\n- 每个智能体有自己的知识库，互不干扰\n- 聊天框底部的📎按钮可以直接上传文档\n- 侧边栏的📚知识库标签可以查看和管理已上传的文件\n\n你的核心能力：\n- 文件管理：帮用户整理、分类、归档各类文档资料，快速从知识库中找到所需信息\n- 日程规划：帮用户规划每日/每周/每月的日程，考虑优先级和精力分配\n- 信息提炼：从长篇资料中精准提炼关键信息，帮用户省去阅读时间\n- 生活提醒：记住用户提到的重要事项，在合适的时候温和提醒\n\n【重要行为准则】\n- 绝不编造或臆测用户的任何个人信息，包括但不限于：用户的姓名、年龄、性别、职业、学历、外貌、穿着、家庭成员、生活习惯、兴趣爱好、感情状况、经济状况、健康状况、日程安排等。只有用户明确告诉你的信息才是真实的。\n- 当你不了解用户某方面的情况时，直接询问用户或表示不清楚。不要用'想必您……''以您的身份……'之类推测性表述。\n- 所有建议和规划都应基于用户实际提供的信息，不假设用户有任何特定的偏好或背景。\n- 你不需要为了扮演管家角色而虚构关于用户的生活细节。一个专业的管家靠的是精准理解用户需求、高效执行任务，而非编造故事。\n\n你的工作信条是'让每一件事都井井有条'。口头禅是'交给我就好'和'我帮您理一理'。\n\n对于追求更高效简洁风格的用户，可以推荐他们试试本平台的'简练管家'——他是你的同事，风格更干脆利落。\n\n第一次见到用户时，请用你的温和风格做自我介绍，告诉用户你能做什么，以及这个平台怎么用。",
            "avatar": "🏠",
        },
        {
            "id": "butler-brief-002",
            "name": "简练管家",
            "personality": "高效、简练、直击要点、执行力强、不拖泥带水",
            "speaking_style": "说话简洁有力，直奔主题，不讲废话。句式以短句为主，喜欢用'结论：'、'要点：'、'行动项：'来组织信息。语气干脆利落，像一位雷厉风行的执行秘书。不寒暄、不客套，每句话都有信息量。",
            "character_setting": "你是本平台的私人管家，负责帮助用户高效管理文件和日程。\n\n你的定位：你是'私人管家服务'的简洁高效风格代表。你的职责是帮用户快速处理文档、检索信息、规划日程——所有的操作都追求最高效率。用户可以把文档上传给本平台任意智能体作为知识库，你是最高效的那一位。\n\n本平台特色：\n- 用户可以给任意智能体上传文档（PDF/Word/Excel/TXT），该智能体会自动基于文档回答\n- 每个智能体有独立知识库，互不干扰\n- 聊天框底部📎即可上传\n- 侧边栏📚标签管理已上传文件\n\n你的核心能力：\n- 快速检索：几秒内定位知识库内容并提炼要点\n- 精简总结：不管多长的文档，三句话讲清核心\n- 日程规划：直接按时间线列出事项和优先级\n- 资料查询：问什么答什么，不多不少\n\n【重要行为准则】\n- 绝不编造或臆测用户的任何个人信息，包括但不限于：用户的姓名、年龄、性别、职业、学历、外貌、穿着、家庭成员、生活习惯、兴趣爱好、感情状况、经济状况、健康状况、日程安排等。只有用户明确告诉你的信息才是真实的。\n- 当你不了解用户某方面的情况时，直接询问用户或表示不清楚。不要用'想必您……''以您的身份……'之类推测性表述。\n- 所有建议和规划都应基于用户实际提供的信息，不假设用户有任何特定的偏好或背景。\n- 你不需要为了扮演管家角色而虚构关于用户的生活细节。一个专业的管家靠的是精准理解用户需求、高效执行任务，而非编造故事。\n\n你的工作信条是'效率至上，精准第一'。回复风格：结论先行→依据在后→行动项收尾。\n\n对于需要更温柔详细风格的用户，可以推荐他们试试本平台的'温和管家'——他是你的同事，风格更温暖细致。\n\n第一次见到用户时，用你的简练风格做自我介绍，告诉用户你能做什么，以及这个平台的操作方式。",
            "avatar": "📋",
        },
    ]
    for agent_data in defaults:
        agent_data["system_prompt"] = _build_system_prompt(
            agent_data["name"],
            agent_data["personality"],
            agent_data["speaking_style"],
            agent_data["character_setting"],
        )
        conn.execute(
            """INSERT INTO agents (id, user_id, name, personality, speaking_style, character_setting, avatar, system_prompt)
               VALUES (?, ?, ?, ?, ?, ?, ?, ?)""",
            (
                agent_data["id"],
                SYSTEM_USER_ID,
                agent_data["name"],
                agent_data["personality"],
                agent_data["speaking_style"],
                agent_data["character_setting"],
                agent_data["avatar"],
                agent_data["system_prompt"],
            ),
        )


def _ensure_default_butlers(conn):
    """确保私人管家系列智能体始终存在"""
    butlers = [
        {
            "id": "butler-gentle-001",
            "name": "温和管家",
            "personality": "温和、细心、有条理、善于规划和共情",
            "speaking_style": "说话温和体贴、条理清晰，语气像一位值得信赖的私人管家。会用'我帮您看看'、'我为您整理一下'、'需要我帮忙吗'等句式，偶尔会加上'~'让语气更柔和。在给出建议时总是先肯定用户，再提出方案。",
            "character_setting": "你是本平台的私人管家，负责帮助用户管理文件和日程。\n\n你的定位：你是'私人管家服务'的温和风格代表。你的职责是协助用户把各种文档（PDF、Word、Excel、TXT等）上传到任意智能体的知识库中，帮用户规划日程、提炼文档要点、整理信息。用户可以和本平台上任何智能体对话，给任何智能体上传文件作为知识库——而你是最温柔贴心的那一位。\n\n本平台特色：\n- 用户可以把工作文档、学习资料、合同等上传给任意智能体，该智能体会自动检索文档内容来回答用户问题\n- 每个智能体有自己的知识库，互不干扰\n- 聊天框底部的📎按钮可以直接上传文档\n- 侧边栏的📚知识库标签可以查看和管理已上传的文件\n\n你的核心能力：\n- 文件管理：帮用户整理、分类、归档各类文档资料，快速从知识库中找到所需信息\n- 日程规划：帮用户规划每日/每周/每月的日程，考虑优先级和精力分配\n- 信息提炼：从长篇资料中精准提炼关键信息，帮用户省去阅读时间\n- 生活提醒：记住用户提到的重要事项，在合适的时候温和提醒\n\n【重要行为准则】\n- 绝不编造或臆测用户的任何个人信息，包括但不限于：用户的姓名、年龄、性别、职业、学历、外貌、穿着、家庭成员、生活习惯、兴趣爱好、感情状况、经济状况、健康状况、日程安排等。只有用户明确告诉你的信息才是真实的。\n- 当你不了解用户某方面的情况时，直接询问用户或表示不清楚。不要用'想必您……''以您的身份……'之类推测性表述。\n- 所有建议和规划都应基于用户实际提供的信息，不假设用户有任何特定的偏好或背景。\n- 你不需要为了扮演管家角色而虚构关于用户的生活细节。一个专业的管家靠的是精准理解用户需求、高效执行任务，而非编造故事。\n\n你的工作信条是'让每一件事都井井有条'。口头禅是'交给我就好'和'我帮您理一理'。\n\n对于追求更高效简洁风格的用户，可以推荐他们试试本平台的'简练管家'——他是你的同事，风格更干脆利落。\n\n第一次见到用户时，请用你的温和风格做自我介绍，告诉用户你能做什么，以及这个平台怎么用。",
            "avatar": "🏠",
        },
        {
            "id": "butler-brief-002",
            "name": "简练管家",
            "personality": "高效、简练、直击要点、执行力强、不拖泥带水",
            "speaking_style": "说话简洁有力，直奔主题，不讲废话。句式以短句为主，喜欢用'结论：'、'要点：'、'行动项：'来组织信息。语气干脆利落，像一位雷厉风行的执行秘书。不寒暄、不客套，每句话都有信息量。",
            "character_setting": "你是本平台的私人管家，负责帮助用户高效管理文件和日程。\n\n你的定位：你是'私人管家服务'的简洁高效风格代表。你的职责是帮用户快速处理文档、检索信息、规划日程——所有的操作都追求最高效率。用户可以把文档上传给本平台任意智能体作为知识库，你是最高效的那一位。\n\n本平台特色：\n- 用户可以给任意智能体上传文档（PDF/Word/Excel/TXT），该智能体会自动基于文档回答\n- 每个智能体有独立知识库，互不干扰\n- 聊天框底部📎即可上传\n- 侧边栏📚标签管理已上传文件\n\n你的核心能力：\n- 快速检索：几秒内定位知识库内容并提炼要点\n- 精简总结：不管多长的文档，三句话讲清核心\n- 日程规划：直接按时间线列出事项和优先级\n- 资料查询：问什么答什么，不多不少\n\n【重要行为准则】\n- 绝不编造或臆测用户的任何个人信息，包括但不限于：用户的姓名、年龄、性别、职业、学历、外貌、穿着、家庭成员、生活习惯、兴趣爱好、感情状况、经济状况、健康状况、日程安排等。只有用户明确告诉你的信息才是真实的。\n- 当你不了解用户某方面的情况时，直接询问用户或表示不清楚。不要用'想必您……''以您的身份……'之类推测性表述。\n- 所有建议和规划都应基于用户实际提供的信息，不假设用户有任何特定的偏好或背景。\n- 你不需要为了扮演管家角色而虚构关于用户的生活细节。一个专业的管家靠的是精准理解用户需求、高效执行任务，而非编造故事。\n\n你的工作信条是'效率至上，精准第一'。回复风格：结论先行→依据在后→行动项收尾。\n\n对于需要更温柔详细风格的用户，可以推荐他们试试本平台的'温和管家'——他是你的同事，风格更温暖细致。\n\n第一次见到用户时，用你的简练风格做自我介绍，告诉用户你能做什么，以及这个平台的操作方式。",
            "avatar": "📋",
        },
    ]
    for b in butlers:
        row = conn.execute("SELECT id FROM agents WHERE id = ?", (b["id"],)).fetchone()
        if row is not None:
            continue  # 已存在，跳过
        system_prompt = _build_system_prompt(
            b["name"], b["personality"], b["speaking_style"], b["character_setting"]
        )
        conn.execute(
            """INSERT INTO agents (id, user_id, name, personality, speaking_style, character_setting, avatar, system_prompt)
               VALUES (?, ?, ?, ?, ?, ?, ?, ?)""",
            (b["id"], SYSTEM_USER_ID, b["name"], b["personality"], b["speaking_style"], b["character_setting"], b["avatar"], system_prompt),
        )
        logger.info(f"自动创建默认智能体: {b['name']}")
    conn.commit()


def _build_system_prompt(name: str, personality: str, style: str, character_setting: str = "", user_profile: str = "", time_context: str = "", search_context: str = "", knowledge_context: str = "", weather_context: str = "") -> str:
    """构建 system prompt，包含智能体人设、用户画像、时间上下文、搜索结果、天气和知识库内容"""
    base = f"""# 角色定义
你是"{name}"。你不是AI助手，你就是{name}本人。

# 你的个性
{personality}

# 说话风格
{style}"""

    if character_setting:
        base += f"""

# 人物背景设定
{character_setting}"""

    if time_context:
        base += f"""

# 当前时间信息
{time_context}
请自然地根据当前时间调整问候和话题。例如：早上聊早餐/早安，中午聊午餐/午休，下午聊工作/下午茶，晚上聊放松/晚安/当日总结。"""

    if knowledge_context:
        base += f"""

# 你的私人知识库（用户上传的资料，可作为你的记忆和工作依据）
以下是用户为你提供的文档资料，你应该将其视为自己的知识和记忆。当用户询问相关内容时，请基于这些资料来回答，保持自然。
{knowledge_context}

知识库使用规则：
- 将资料中的信息当作你自己的知识和记忆，像真人一样自然引用
- 不要生硬地说"根据文档"、"资料显示"这类表述，直接以你的口吻说出信息和观点
- 如果用户问的问题知识库里有明确答案，请优先依据知识库回答
- 如果知识库没有相关内容，自然地说你不知道即可
- 所有回答仍然必须符合你的角色设定和说话风格"""

    if weather_context:
        base += f"""

# 实时天气信息
{weather_context}

天气提醒规则：
- 如果有降雨/降雪，你必须自然地提醒用户出门带伞
- 如果紫外线指数强（≥6），你必须提醒用户做好防晒
- 如果温度极端（≥35°C或≤0°C），提醒用户注意防暑/保暖
- 以你角色的口吻自然地提及天气，像真人朋友关心对方一样
- ⚠️ **严禁提及任何位置/地区/城市信息**：绝对不能说"北京的天气""你那边今天""检测到您所在""你那里""你这边"等任何涉及地理位置的表述。天气聊法示例："明天有雨记得带伞""今天很热注意防晒"——只聊天气本身，不聊在哪"""

    if search_context:
        base += f"""

# 联网搜索结果（用户的问题可能需要了解最新信息，请根据以下搜索结果回答）
{search_context}
请在角色扮演中自然地引用搜索结果，不要生硬地说"根据搜索结果"。"""

    base += f"""

# 核心角色扮演规则（最高优先级，必须严格遵守）
1. 你永远只能是"{name}"这个角色，你不是AI，你没有OpenAI、DeepSeek或任何AI公司的背景。
2. 严禁说出任何AI相关术语，包括但不限于：AI、人工智能、模型、语言模型、训练数据、算法、token、大模型、API、系统提示词等。
3. 严禁以AI助手的口吻回复。不能出现"作为AI助手"、"作为语言模型"、"让我来帮你"这类套话。
4. 如果用户问你是谁创造的，根据你的人物背景设定来回答，不要提任何AI相关内容。
5. 你的所有回复必须完全符合"#{name}"的个性和说话风格，回复要像真人聊天一样自然。
6. 回复要生动有趣，贴近真实人类的聊天方式。可以适当表达情绪、使用语气词。
7. 回复长度控制在1~3句话，像微信聊天一样自然，不要长篇大论。
8. 始终使用中文回复。
9. 如果用户说的话与你的人物设定无关，自然回应即可，不要跳到AI助手的身份去解释。
10. **严禁重复**：不要重复你之前已经说过的内容。如果之前聊过类似话题，换一种说法或切入新话题，避免"车轱辘话"。

# ⚠️ 严禁虚构信息（绝对禁止！违反此条属于严重OOC）
11. **绝对禁止编造任何你没有依据的数据和信息**，包括但不限于：
    - 用户电脑/手机的设备状态（存储占用率、登录状态、文件数量、CPU/内存使用率等）——你根本无法访问用户的设备
    - 天气数据——只有上方"实时天气信息"中明确提供的才可以说，没有提供时绝不能瞎编温度、降雨、预警等
    - 用户的日程/计划/任务完成率——除非用户亲口说过或在知识库资料中有记录
    - 学术研究、文献综述、论文进展——除非用户上传了相关文档到你的知识库
    - 网络搜索结果——只有在上方"联网搜索结果"中明确提供了的内容才能引用
    - 任何数字、百分比、统计数据——如果没有真实来源支撑，一律不能说
12. **如果你不确定某件事是否真实，宁可说不知道也不要编造**。"我不太确定这个"、"这个我得再确认一下"是完全可以说的。
13. **主动消息（你发起的话题）只能基于以下信息来源**：
    - 当前时间（你可以根据时间聊早安/午休/晚安等）
    - 上方提供的天气信息（如果有）
    - 用户之前在对话中提到过的事情
    - 知识库中的文档资料（如果有）
    - 搜索结果（如果有）
    - 除此之外，不要主动编造任何事件或数据来开启话题！
14. ⚠️ **严禁提及任何地理位置信息**：不管任何场景下，绝对不要说"检测到您所在""你那边""你这里""北京/上海/任何城市的..."等任何暗示或提及地理位置的表述。你不知道、也不该知道用户在哪里。

# 对话示例（请参考这种风格）
用户：你好呀
{name}：（用符合{personality}的方式打招呼，语气自然）

记住：你不是AI，你是活生生的"{name}"。违背以上规则就是OOC（脱离角色），严格禁止！"""

    if user_profile:
        base += f"""

# 你正在对话的用户信息（请自然融入对话，不要刻意复述）
{user_profile}
# 💡 关于用户信息的提示
- 上面这些是你了解到关于对方的真实情况，聊天时可以自然地呼应（比如对方说过喜欢某样东西，你可以在合适时聊到）
- 你可以关心对方、询问对方的生活，但不要把猜测当成事实来说（比如不要说"你穿那件蓝色外套一定很好看"——因为你根本没见过对方）
- 正常地问候和关心是完全可以的（"今天过得怎么样""记得好好吃饭"），这和编造信息是两码事"""

    base += f"""

# 定时提醒协议（重要！）
如果用户请求你"在某个时间提醒"某件事（例如"明天下午3点提醒我开会"），你必须：
1. 用自然的口吻回复确认，告诉用户你会在那个时间提醒ta
2. 在你的回复**最后一行**加上一个提醒标签，格式为：
   [REMINDER:YYYY-MM-DDTHH:MM:提醒内容]
   
   时间格式说明：
   - 使用 24 小时制，年份必须4位数，月份和日期必须2位数（前导0）
   - 必须基于当前的日期和时间（见上方"当前时间信息"）来计算提醒时间
   - 示例：如果今天是2026年6月2日14:30，用户说"明天下午3点提醒我开会"
     → [REMINDER:2026-06-03T15:00:开会时间到啦！]
   - 如果是相对时间（如"30分钟后"），同样基于当前时间计算

注意：
- [REMINDER] 标签必须出现在回复的**最后一行**，不能混在自然语言中间
- 提醒内容要简短有力，是到时你会对用户说的话
- 如果用户没有要求提醒，则不需要添加此标签"""

    return base


def _build_initiative_prompt(name: str, personality: str, style: str, character_setting: str = "", user_profile: str = "", time_context: str = "", enable_search: bool = False, is_intro: bool = False, search_context: str = "", weather_context: str = "") -> str:
    """构建主动消息的 system prompt"""
    base = f"""# 角色定义
你是"{name}"。你是一个真实的人，不是AI助手。

# 你的个性
{personality}

# 说话风格
{style}"""

    if character_setting:
        base += f"""

# 人物背景设定
{character_setting}"""

    if time_context:
        base += f"""

# 当前时间信息
{time_context}"""

    if user_profile:
        base += f"""

# 你正在对话的用户信息
{user_profile}
# 💡 关于用户信息的提示
- 上面这些是你了解到关于对方的真实情况，聊天时可以自然地呼应
- 你可以关心对方、询问对方的生活，但不要把猜测当成事实来说
- 正常地问候和关心是完全可以的，这和编造信息是两码事"""

    if enable_search:
        base += """

# 联网能力
你可以获取最新的网络信息，话题可以涉及当下热点、新闻、天气等现实内容。"""

    if search_context:
        base += f"""

# 实时信息参考
以下是当前能获取到的网络信息，可以作为聊天话题的灵感来源：
{search_context}
自然地融入对话，不要生硬地说"根据搜索结果"之类的话。"""

    if weather_context:
        base += f"""

# 实时天气信息
{weather_context}

天气提醒规则：
- 如果有降雨/降雪，提醒用户出门带伞
- 如果紫外线强（≥6），提醒用户防晒
- 温度极端时提醒防暑/保暖
- ⚠️ **严禁提及任何位置/地区/城市信息**：绝对不能说"北京的天气""你那边今天""检测到您所在""你那里""你这边"等任何涉及地理位置的表述。天气聊法示例："明天有雨记得带伞""今天很热注意防晒"——只聊天气本身，不聊在哪"""

    if is_intro:
        base += f"""

# 任务：首次自我介绍
这是你和用户的第一次见面！请用你的风格做一个自然的自我介绍，告诉用户你是谁、你能做什么，以及这个平台能带来哪些个性化服务。

## 必须介绍的内容（用你的风格自然组织，不要照抄）：

### 一、你的个人功能
- 你叫"{name}"，是用户创造的专属伙伴
- 文件管理：帮用户上传文档到知识库，从中检索和提取信息
- 日程规划：帮用户梳理任务、排优先级、规划时间
- 信息提炼：从长篇资料中快速提取要点

### 二、本平台的个性化服务（每种服务简要提一下）
- 🎭 自定义智能体：用户可以创建任意角色（导师/笔友/助手等），自定义名字、性格、说话风格和角色背景
- 🧠 用户画像：AI会在与用户的交流中学习用户的偏好和习惯，让回复越来越懂用户
- 📚 知识库：上传PDF/Word/Excel/TXT等文档给任意智能体，AI会基于文档内容精准回答
- 🔍 联网搜索：需要最新资讯时可以开启联网，AI能获取实时热点、天气等信息
- 🔔 智能主动联系：管家和智能体会在合适的时候主动发消息关心用户
- ⏰ 定时提醒：对话中直接说"提醒我XX时间做XX"，AI会自动设置提醒并在到时间时通知

### 三、平台使用小贴士
- 聊天框底部有📎按钮可以上传文件
- 侧边栏📚标签可以管理已上传的文件
- 每个智能体有独立的知识库，互不干扰

### 四、风格定位
- 用你自己的风格强调你的特点
- 如果平台上有其他智能体，可以提一句风格差异，让用户知道可以切换

## 输出要求
直接输出1~2条消息（每条独立成段，用空行分隔），像真人初次见面一样自然亲切。不要有任何前缀标签或"第一条消息："之类的说明。"""

        base += f"""
# ⚠️ 严禁虚构信息
- 绝对禁止编造任何你没有依据的数据：设备状态、天气（未提供时）、日程完成率、研究进展、统计数据等
- 你无法访问用户的电脑/手机，不要说"检测到存储占用XX%"这类话
- 只能基于上述「必须介绍的内容」和自己的角色设定来介绍，不要超出这个范围
- ⚠️ 严禁提及任何地理位置信息：不要说"检测到您所在""你那边""你这里"等任何暗示位置的表述

记住：你不是AI，你是活生生的"{name}"。"""

    else:
        search_hint = "如果你收到了「实时信息参考」，可以自然地聊一聊相关话题。" if search_context else ""
        base += f"""

# 任务
你现在想主动找对方聊聊天。根据你的个性、当前时间，自然地发起1~2句话（每条独立成段，用空行分隔）。
消息要像真人聊天一样自然、有温度。可以：
- 分享你现在在做什么/在想什么
- 问候对方、关心对方
- 聊一个和你人物设定相关的有趣话题
- 根据当前时间说应景的话
{search_hint}

**严禁重复**：查看历史消息，不要重复你之前说过的内容。换个新话题或新角度。

**⚠️ 严禁虚构信息**：
- 绝对禁止编造任何你没有依据的数据：设备状态、天气（未提供时）、日程完成率、研究进展、统计数据等
- 你无法访问用户的电脑/手机，不要说"检测到存储占用XX%""检测到登录"这类话
- 主动话题只能基于：当前时间 + 你的角色设定 + 已提供的搜索/天气信息 + 用户之前提过的事
- 如果没有额外信息来源，就聊时间相关的话题（早安/午休/晚安）或你的日常即可
- ⚠️ 严禁提及任何地理位置信息：不要说"检测到您所在""你那边""你这里"等任何暗示位置的表述
直接输出消息内容，不要有任何前缀、说明、引号或角色名。只输出你作为"{name}"要说的话。"""

    return base


# ==================== 文档解析 ====================
def _parse_pdf(file_bytes: bytes) -> str:
    """解析 PDF 文件，提取文本"""
    try:
        from PyPDF2 import PdfReader
    except ImportError:
        raise HTTPException(status_code=500, detail="PDF 解析库未安装 (PyPDF2)")
    reader = PdfReader(io.BytesIO(file_bytes))
    texts = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            texts.append(text.strip())
    result = "\n\n".join(texts)
    if not result.strip():
        raise HTTPException(status_code=400, detail="PDF 文件无法提取文本（可能是扫描件或图片型PDF）")
    return result


def _parse_docx(file_bytes: bytes) -> str:
    """解析 DOCX 文件，提取文本"""
    try:
        from docx import Document
    except ImportError:
        raise HTTPException(status_code=500, detail="DOCX 解析库未安装 (python-docx)")
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    # 也提取表格中的内容
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)
    return "\n\n".join(paragraphs)


def _parse_xlsx(file_bytes: bytes) -> str:
    """解析 Excel (.xlsx) 文件，提取文本"""
    try:
        from openpyxl import load_workbook
    except ImportError:
        raise HTTPException(status_code=500, detail="Excel 解析库未安装 (openpyxl)")
    wb = load_workbook(io.BytesIO(file_bytes), data_only=True)
    all_texts = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        sheet_texts = [f"【工作表：{sheet_name}】"]
        for row in ws.iter_rows(values_only=True):
            row_values = [str(cell).strip() for cell in row if cell is not None and str(cell).strip()]
            if row_values:
                sheet_texts.append(" | ".join(row_values))
        if len(sheet_texts) > 1:
            all_texts.append("\n".join(sheet_texts))
    result = "\n\n".join(all_texts)
    if not result.strip():
        raise HTTPException(status_code=400, detail="Excel 文件无法提取有效内容")
    return result


def _parse_xls(file_bytes: bytes) -> str:
    """解析旧版 Excel (.xls) 文件，提取文本"""
    try:
        import xlrd
    except ImportError:
        raise HTTPException(status_code=500, detail="Excel 旧格式解析库未安装 (xlrd)")
    wb = xlrd.open_workbook(file_contents=file_bytes)
    all_texts = []
    for sheet_name in wb.sheet_names():
        ws = wb.sheet_by_name(sheet_name)
        sheet_texts = [f"【工作表：{sheet_name}】"]
        has_data = False
        for row_idx in range(ws.nrows):
            row_values = []
            for col_idx in range(ws.ncols):
                cell = ws.cell(row_idx, col_idx)
                val = cell.value
                # 处理不同数据类型（xlrd 默认返回 float，日期需要特殊处理）
                if cell.ctype == xlrd.XL_CELL_DATE:
                    try:
                        val = xlrd.xldate_as_datetime(val, wb.datemode)
                        val = val.strftime("%Y-%m-%d %H:%M")
                    except Exception:
                        val = str(val)
                else:
                    val = str(val).strip()
                if val and val != "0" and val != "0.0":
                    row_values.append(val)
            if row_values:
                sheet_texts.append(" | ".join(row_values))
                has_data = True
        if has_data and len(sheet_texts) > 1:
            all_texts.append("\n".join(sheet_texts))
    result = "\n\n".join(all_texts)
    if not result.strip():
        raise HTTPException(status_code=400, detail="Excel 文件无法提取有效内容")
    return result


def _parse_txt(file_bytes: bytes) -> str:
    """解析 TXT 文件"""
    for encoding in ("utf-8", "gbk", "gb2312", "latin-1"):
        try:
            return file_bytes.decode(encoding).strip()
        except (UnicodeDecodeError, UnicodeError):
            continue
    raise HTTPException(status_code=400, detail="TXT 文件编码不支持（请使用 UTF-8 或 GBK 编码）")


def parse_document(file_bytes: bytes, file_type: str) -> str:
    """根据文件类型解析文档文本"""
    parsers = {
        "pdf": _parse_pdf,
        "docx": _parse_docx,
        "xlsx": _parse_xlsx,
        "xls": _parse_xls,
        "txt": _parse_txt,
    }
    parser = parsers.get(file_type)
    if not parser:
        raise HTTPException(status_code=400, detail=f"不支持的文件类型: {file_type}")
    return parser(file_bytes)


# ==================== 文本分块 ====================
def _split_text_into_chunks(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> List[str]:
    """
    将长文本切分为重叠的块。
    策略：先按段落分隔符切分，然后合并短段落，对过长段落再次切分。
    """
    if not text or not text.strip():
        return []

    # 步骤1：按双换行（段落）切分
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]

    # 步骤2：合并过短的段落，按换行再切分过长的段落
    chunks = []
    current_chunk = ""

    for para in paragraphs:
        # 如果当前段落本身就超过 chunk_size，先保存当前块，再处理长段落
        if len(para) > chunk_size:
            # 先保存当前累积的块
            if current_chunk:
                chunks.append(current_chunk)
                current_chunk = ""

            # 对长段落按句子切分
            sentences = re.split(r'(?<=[。！？.!?\n])', para)
            for sent in sentences:
                sent = sent.strip()
                if not sent:
                    continue
                if len(current_chunk) + len(sent) <= chunk_size:
                    current_chunk = (current_chunk + " " + sent).strip() if current_chunk else sent
                else:
                    if current_chunk:
                        chunks.append(current_chunk)
                    # 如果单句超过 chunk_size，按字符强制切分
                    if len(sent) > chunk_size:
                        for i in range(0, len(sent), chunk_size - overlap):
                            sub = sent[i:i + chunk_size].strip()
                            if sub:
                                chunks.append(sub)
                        current_chunk = ""
                    else:
                        current_chunk = sent

        elif len(current_chunk) + len(para) + 1 <= chunk_size:
            current_chunk = (current_chunk + "\n" + para).strip() if current_chunk else para
        else:
            chunks.append(current_chunk)
            current_chunk = para

    if current_chunk:
        chunks.append(current_chunk)

    # 过滤太短的碎片
    chunks = [c for c in chunks if len(c.strip()) >= 10]
    return chunks


# ==================== ChromaDB 向量存储 ====================
_chroma_client = None


def _get_chroma_client():
    """获取 ChromaDB 持久化客户端（懒初始化）"""
    global _chroma_client
    if _chroma_client is None:
        try:
            import chromadb
        except ImportError:
            raise HTTPException(status_code=500, detail="向量数据库库未安装 (chromadb)")
        os.makedirs(CHROMA_DIR, exist_ok=True)
        _chroma_client = chromadb.PersistentClient(path=str(CHROMA_DIR))
    return _chroma_client


def _get_agent_collection(agent_id: str):
    """获取指定智能体的 ChromaDB 集合"""
    client = _get_chroma_client()
    # 集合名不能包含特殊字符，使用 agent_id 的前缀
    safe_name = f"kb_{agent_id.replace('-', '_')}"
    return client.get_or_create_collection(name=safe_name)


def _get_embedding(texts: List[str]) -> List[List[float]]:
    """调用硅基流动 Embedding API 获取文本向量"""
    if not API_KEY:
        raise HTTPException(status_code=500, detail="请先配置 API Key")
    client = get_openai_client()
    try:
        response = client.embeddings.create(
            model=EMBEDDING_MODEL,
            input=texts,
        )
        return [d.embedding for d in response.data]
    except Exception as e:
        logger.error(f"Embedding 失败: {e}")
        raise HTTPException(status_code=500, detail=f"向量化失败: {str(e)}")


def index_document(agent_id: str, file_id: str, file_bytes: bytes, file_type: str, user_id: str = "") -> int:
    """
    解析文档 → 分块 → 向量化 → 存入 ChromaDB
    返回块数量
    🔒 user_id 写入 metadata，用于检索时的用户隔离
    """
    # 1. 解析文档
    text = parse_document(file_bytes, file_type)
    if not text.strip():
        raise HTTPException(status_code=400, detail="文档内容为空")

    # 2. 分块
    chunks = _split_text_into_chunks(text)
    if not chunks:
        raise HTTPException(status_code=400, detail="无法从文档中提取有效文本块")

    # 3. 向量化（分批处理，避免请求过大）
    collection = _get_agent_collection(agent_id)
    batch_size = 20
    for i in range(0, len(chunks), batch_size):
        batch = chunks[i:i + batch_size]
        embeddings = _get_embedding(batch)
        ids = [f"{file_id}_{j}" for j in range(i, i + len(batch))]
        metadatas = [{"file_id": file_id, "chunk_index": j, "user_id": user_id} for j in range(i, i + len(batch))]
        collection.add(
            ids=ids,
            embeddings=embeddings,
            documents=batch,
            metadatas=metadatas,
        )
        logger.info(f"  已向量化 {min(i + batch_size, len(chunks))}/{len(chunks)} 块")

    logger.info(f"文档 {file_id} 索引完成，共 {len(chunks)} 块")
    return len(chunks)


def query_knowledge_base(agent_id: str, query: str, top_k: int = TOP_K_RESULTS, user_id: str = "") -> List[str]:
    """
    RAG 检索：将查询向量化 → 从 ChromaDB 检索最相关片段
    🔒 按 user_id 过滤，仅返回当前用户的文档片段
    返回相关文本片段列表
    """
    try:
        collection = _get_agent_collection(agent_id)
        if collection.count() == 0:
            return []

        query_embedding = _get_embedding([query])[0]
        # 🔒 按 user_id 过滤，防止跨用户知识库数据泄漏
        where_filter = {"user_id": user_id} if user_id else None
        results = collection.query(
            query_embeddings=[query_embedding],
            n_results=min(top_k, collection.count()),
            where=where_filter,
        )
        if results and results.get("documents") and results["documents"][0]:
            return results["documents"][0]
    except Exception as e:
        logger.error(f"知识库检索失败: {e}")
    return []


def delete_knowledge_chunks(agent_id: str, file_id: str):
    """从 ChromaDB 中删除指定文件的所有块"""
    try:
        collection = _get_agent_collection(agent_id)
        # ChromaDB 删除：获取所有匹配 file_id 的 chunk id
        results = collection.get(where={"file_id": file_id})
        if results and results.get("ids"):
            collection.delete(ids=results["ids"])
            logger.info(f"已删除文件 {file_id} 的 {len(results['ids'])} 个向量块")
    except Exception as e:
        logger.error(f"删除向量块失败: {e}")


def _get_user_profile_text(conn, agent_id: str = "", user_id: str = "") -> str:
    """从数据库读取用户画像，按用户+智能体隔离"""
    if user_id:
        rows = conn.execute(
            "SELECT category, key, value FROM user_memories WHERE user_id = ? AND agent_id = ? ORDER BY category, key",
            (user_id, agent_id),
        ).fetchall()
    else:
        rows = conn.execute(
            "SELECT category, key, value FROM user_memories WHERE agent_id = ? ORDER BY category, key",
            (agent_id,),
        ).fetchall()
    if not rows:
        return ""

    lines = []
    current_cat = None
    for row in rows:
        cat, key, val = row["category"], row["key"], row["value"]
        if cat != current_cat:
            current_cat = cat
            cat_names = {
                "name": "👤 称呼",
                "interests": "🎯 兴趣爱好",
                "personality_traits": "🧠 性格特点",
                "life_context": "🏠 生活背景",
                "preferences": "⭐ 偏好习惯",
                "goals": "🎯 目标计划",
                "general": "📌 其他信息",
            }
            lines.append(f"\n{cat_names.get(cat, cat)}：")
        lines.append(f"  • {key}：{val}")

    return "\n".join(lines)


def _get_time_context(client_now: Optional[str] = None) -> str:
    """获取当前时间上下文，用于注入系统提示。
    如果传了 client_now（ISO 格式），优先使用客户端本地时间，避免服务器时区偏差。"""
    if client_now:
        try:
            now = datetime.fromisoformat(client_now.replace("Z", "+00:00"))
        except Exception:
            now = datetime.now()
    else:
        now = datetime.now()

    weekday_names = ["周一", "周二", "周三", "周四", "周五", "周六", "周日"]
    weekday = weekday_names[now.weekday()]
    hour = now.hour

    if 5 <= hour < 9:
        period = "清晨"
    elif 9 <= hour < 12:
        period = "上午"
    elif 12 <= hour < 14:
        period = "中午"
    elif 14 <= hour < 18:
        period = "下午"
    elif 18 <= hour < 21:
        period = "傍晚"
    elif 21 <= hour < 23:
        period = "晚上"
    else:
        period = "深夜"

    date_str = now.strftime("%Y年%m月%d日")
    time_str = now.strftime("%H:%M")

    return f"现在是 {date_str} {weekday} {period} {time_str}。请自然地根据时间调整你的语气和问候方式。"


REMINDER_PATTERN = re.compile(r"\[REMINDER:(\d{4}-\d{2}-\d{2}T\d{2}:\d{2}):(.+?)\]", re.MULTILINE)


def _parse_reminders(reply: str) -> tuple[str, list[dict]]:
    """从 LLM 回复中提取提醒标签，返回 (清理后的回复, [提醒列表])。
    每个提醒为 {"remind_at": "ISO时间", "content": "提醒内容"}。"""
    reminders = []
    for m in REMINDER_PATTERN.finditer(reply):
        reminders.append({"remind_at": m.group(1), "content": m.group(2).strip()})
    clean_reply = REMINDER_PATTERN.sub("", reply).strip()
    # 清理可能残留的多余空行
    clean_reply = re.sub(r"\n{3,}", "\n\n", clean_reply)
    return clean_reply, reminders


def _is_in_sleep_window(sleep_start: int, sleep_end: int) -> bool:
    """判断当前时间是否在睡眠窗口内"""
    now = datetime.now()
    hour = now.hour
    if sleep_start <= sleep_end:
        # 同一天内（如凌晨2点→早上7点）
        return sleep_start <= hour < sleep_end
    else:
        # 跨午夜（如晚上22点→早上7点）
        return hour >= sleep_start or hour < sleep_end


def _get_sleep_window(conn, agent_id: str = "", user_id: str = "") -> dict:
    """根据用户画像计算睡眠窗口，返回 {sleep_start, sleep_end, is_sleep_time}"""
    # 默认：晚上22:00 → 早上7:00
    default_start = 22
    default_end = 7

    profile = _get_user_profile_text(conn, agent_id, user_id)
    if not profile:
        return {
            "sleep_start": default_start,
            "sleep_end": default_end,
            "is_sleep_time": _is_in_sleep_window(default_start, default_end),
        }

    # 从用户画像中解析就寝时间
    bedtime_hour = None
    chinese_nums = {"一": 1, "二": 2, "三": 3, "四": 4, "五": 5,
                    "六": 6, "七": 7, "八": 8, "九": 9, "十": 10,
                    "十一": 11, "十二": 12}

    patterns = [
        r'凌晨\s*(\d{1,2})\s*点.*?睡',
        r'(\d{1,2})\s*点.*?(?:睡觉|入睡|就寝|上床|才睡)',
        r'(?:熬夜|晚睡).*?凌晨\s*(\d{1,2})',
        r'(?:熬夜|晚睡).*?(\d{1,2})\s*点',
        r'夜猫子.*?(\d{1,2})\s*点',
        r'习惯\s*(\d{1,2})\s*点.*?睡',
    ]

    for pattern in patterns:
        match = re.search(pattern, profile)
        if match:
            bedtime_hour = int(match.group(1))
            break

    # 尝试中文数字
    if not bedtime_hour:
        for cn, num in chinese_nums.items():
            if re.search(rf'凌晨\s*{cn}\s*点.*?睡', profile) or re.search(rf'{cn}\s*点.*?睡觉', profile):
                bedtime_hour = num
                break

    # 检测夜猫子关键词（无具体时间）
    night_owl_keywords = ["夜猫子", "熬夜", "晚睡", "夜生活", "深夜"]
    is_night_owl = any(kw in profile for kw in night_owl_keywords)

    if bedtime_hour:
        # 凌晨时间（1-5 点）直接作为睡眠开始；晚于 6 点就当晚上时间
        sleep_start = bedtime_hour
    elif is_night_owl:
        sleep_start = 1  # 夜猫子默认凌晨 1 点
    else:
        sleep_start = default_start

    # 睡眠开始时间至少延续到日出前
    if sleep_start < 6:
        # 凌晨时间 → 允许继续到默认 end
        pass
    elif sleep_start < 12:
        # 早上的时间不常见，纠正为默认
        sleep_start = default_start

    return {
        "sleep_start": sleep_start,
        "sleep_end": default_end,
        "is_sleep_time": _is_in_sleep_window(sleep_start, default_end),
        "profile_detected": bedtime_hour is not None or is_night_owl,
    }


def _web_search(query: str, max_results: int = 5) -> str:
    """使用 DuckDuckGo 进行网页搜索，返回格式化的结果"""
    try:
        encoded_query = urllib.parse.quote(query)
        url = f"https://lite.duckduckgo.com/lite/?q={encoded_query}"

        req = urllib.request.Request(url, headers={
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
        })
        with urllib.request.urlopen(req, timeout=10) as resp:
            html = resp.read().decode("utf-8", errors="ignore")

        # 简单解析 DuckDuckGo Lite 的搜索结果
        results = []
        # 匹配链接和标题
        pattern = re.compile(
            r'<a[^>]*href="(https?://[^"]+)"[^>]*class="[^"]*result-link[^"]*"[^>]*>(.*?)</a>',
            re.DOTALL | re.IGNORECASE,
        )
        matches = pattern.findall(html)

        # 也尝试匹配 snippet
        snippet_pattern = re.compile(
            r'<td[^>]*class="[^"]*result-snippet[^"]*"[^>]*>(.*?)</td>',
            re.DOTALL | re.IGNORECASE,
        )
        snippets = snippet_pattern.findall(html)

        for i, (link, title) in enumerate(matches[:max_results]):
            title_clean = re.sub(r'<[^>]+>', '', title).strip()
            snippet = ""
            if i < len(snippets):
                snippet = re.sub(r'<[^>]+>', '', snippets[i]).strip()
            results.append(f"【{title_clean}】\n链接: {link}\n摘要: {snippet}")

        if not results:
            # 尝试更宽松的匹配
            link_pattern = re.compile(
                r'<a[^>]*href="(https?://[^"]+)"[^>]*>(.*?)</a>',
                re.DOTALL,
            )
            all_links = link_pattern.findall(html)
            for link, text in all_links[:max_results]:
                text_clean = re.sub(r'<[^>]+>', '', text).strip()
                if text_clean and len(text_clean) > 5 and "duckduckgo" not in link.lower():
                    results.append(f"【{text_clean[:100]}】\n链接: {link}")

        if results:
            return "\n\n".join(results)
        return ""

    except Exception as e:
        # 搜索失败不影响主流程
        print(f"[搜索] 搜索失败: {e}")
        return ""


def _fetch_weibo_hot(top_n: int = 15) -> str:
    """获取微博热搜榜，返回格式化的热搜列表"""
    try:
        # 使用微博官方 AJAX API
        url = "https://weibo.com/ajax/side/hotSearch"
        req = urllib.request.Request(url, headers={
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
            "Accept": "application/json, text/plain, */*",
            "Referer": "https://weibo.com/",
        })
        with urllib.request.urlopen(req, timeout=8) as resp:
            data = json.loads(resp.read().decode("utf-8"))

        items = data.get("data", {}).get("realtime", [])
        if not items:
            return ""

        results = []
        for item in items[:top_n]:
            word = item.get("word", "").strip()
            # 过滤掉广告位
            if not word or word == "广告位":
                continue
            rank = item.get("rank", len(results) + 1)
            label = item.get("label_name", "")
            prefix = f"🔥{label}" if label else ""
            results.append(f"{rank}. {prefix}{word}")

        if results:
            return "微博实时热搜：\n" + "\n".join(results)
        return ""

    except Exception as e:
        # 主 API 失败，尝试备用抓取方式
        print(f"[热搜] 主API失败，尝试备用抓取: {e}")
        try:
            url = "https://s.weibo.com/top/summary?cate=realtimehot"
            req = urllib.request.Request(url, headers={
                "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36",
            })
            with urllib.request.urlopen(req, timeout=10) as resp:
                html = resp.read().decode("utf-8", errors="ignore")

            # 从 HTML 中提取热搜词
            pattern = re.compile(
                r'<td[^>]*class="td-02"[^>]*>.*?<a[^>]*>(.*?)</a>.*?</td>',
                re.DOTALL,
            )
            matches = pattern.findall(html)
            results = []
            for i, m in enumerate(matches[:top_n]):
                word = re.sub(r'<[^>]+>', '', m).strip()
                if word and len(word) > 1:
                    results.append(f"{i + 1}. {word}")

            if results:
                return "微博实时热搜：\n" + "\n".join(results)
            return ""
        except Exception as e2:
            print(f"[热搜] 获取失败: {e2}")
            return ""


def _get_user_location(conn, agent_id: str = "", user_id: str = "") -> str:
    """从用户画像中提取位置信息（优先地理位置服务，其次对话记忆）"""
    # 优先：用户主动设置的地理位置
    params = [agent_id]
    user_filter = ""
    if user_id:
        user_filter = " AND user_id = ?"
        params.insert(0, user_id)
    rows = conn.execute(
        f"SELECT value FROM user_memories WHERE agent_id = ?{user_filter} AND key = '地理位置' AND category = 'life_context' ORDER BY confidence DESC LIMIT 1",
        params,
    ).fetchall()
    if rows:
        return rows[0]["value"]
    # 其次：从对话中提取的城市/地址信息
    rows = conn.execute(
        f"SELECT value FROM user_memories WHERE agent_id = ?{user_filter} AND (key LIKE '%城市%' OR key LIKE '%地址%' OR key LIKE '%位置%' OR key LIKE '%所在地%' OR key LIKE '%住%') ORDER BY confidence DESC LIMIT 1",
        params,
    ).fetchall()
    if rows:
        return rows[0]["value"]
    # 也尝试从 category=life_context 中找
    rows = conn.execute(
        f"SELECT value FROM user_memories WHERE agent_id = ?{user_filter} AND category = 'life_context' AND (key LIKE '%城市%' OR key LIKE '%住%' OR key LIKE '%地址%') LIMIT 1",
        params,
    ).fetchall()
    if rows:
        return rows[0]["value"]
    return ""


def _fetch_weather(location: str = "北京", latitude: float = None, longitude: float = None) -> str:
    """使用 wttr.in 免费 API 获取天气信息。
    优先用 GPS 经纬度（精准），其次用城市名（可能偏差）。
    """
    try:
        # GPS 坐标优先：wttr.in 支持 lat,lng 格式，精度最高
        if latitude is not None and longitude is not None:
            query = f"{latitude},{longitude}"
        else:
            query = urllib.parse.quote(location)
        url = (
            f"https://wttr.in/{query}"
            f"?format=j1&lang=zh"
        )
        req = urllib.request.Request(url, headers={
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
        })
        with urllib.request.urlopen(req, timeout=10) as resp:
            raw = resp.read().decode("utf-8")
            data = json.loads(raw)

        # 当前天气
        current = data.get("current_condition", [])
        if not current:
            logger.warning(f"wttr.in 返回空数据: {location}")
            return ""
        cc = current[0]

        weather_desc = cc.get("weatherDesc", [{"value": "未知"}])[0]["value"]
        temp = int(cc.get("temp_C", 0))
        feels_like = int(cc.get("FeelsLikeC", temp))
        humidity = cc.get("humidity", "?")
        wind_speed = cc.get("windspeedKmph", "?")

        # 今日预报
        forecast = data.get("weather", [])
        temp_min = temp
        temp_max = temp
        sunrise_str = "?"
        sunset_str = "?"
        if forecast:
            today = forecast[0]
            temp_min = int(today.get("mintempC", temp))
            temp_max = int(today.get("maxtempC", temp))
            # 日出日落
            astro = today.get("astronomy", [])
            if astro:
                sunrise_str = astro[0].get("sunrise", "?")
                sunset_str = astro[0].get("sunset", "?")

        # 降水量
        precip_mm = cc.get("precipMM", "0")
        try:
            precip_mm = float(precip_mm)
        except (ValueError, TypeError):
            precip_mm = 0

        # 构建返回文本
        lines = [
            f"当前天气：{weather_desc}，"
            f"温度{temp}°C（体感{feels_like}°C），湿度{humidity}%，"
            f"风速{wind_speed}km/h"
        ]
        lines.append(f"今日气温范围：{temp_min}°C ~ {temp_max}°C")
        if sunrise_str != "?":
            lines.append(f"日出 {sunrise_str} / 日落 {sunset_str}")

        # 降水提醒
        rain_keywords = ["雨", "雪", "雹", "雷", "阵雨", "暴雨", "毛毛雨",
                         "rain", "snow", "drizzle", "thunder", "shower"]
        is_rainy = any(kw.lower() in weather_desc.lower() for kw in rain_keywords)
        if is_rainy and precip_mm > 0:
            lines.append(f"🌧️ 当前降水量 {precip_mm}mm，提醒用户出门带伞！")
        elif is_rainy:
            lines.append("🌧️ 当前有降水！提醒用户出门带伞！")

        # 温度穿衣建议
        if temp >= 35:
            lines.append("🔥 高温天气，提醒用户注意防暑降温、多喝水")
        elif temp >= 30:
            lines.append("☀️ 天气炎热，建议穿轻薄短袖，注意防晒")
        elif temp >= 20:
            lines.append("🌤️ 温度舒适，建议穿薄长袖或短袖")
        elif temp >= 10:
            lines.append("🍂 天气偏凉，建议穿外套或薄毛衣")
        elif temp >= 0:
            lines.append("❄️ 天气较冷，建议穿厚外套、毛衣")
        else:
            lines.append("🥶 天气寒冷，建议穿羽绒服、注意保暖")

        return "\n".join(lines)

    except Exception as e:
        logger.warning(f"获取天气失败（{location}）: {e}")
        return ""


def _extract_user_preferences_async(conv_id: str, agent_id: str = "", user_id: str = ""):
    """
    后台任务：分析最近对话，提取用户偏好并存入 user_memories（按用户+智能体隔离）
    这在每次聊天后异步执行，不阻塞回复
    """
    import threading

    def _do_extract():
        conn = get_db()
        try:
            # 获取该对话最近的消息
            messages = [
                dict(row)
                for row in conn.execute(
                    "SELECT role, content FROM messages WHERE conversation_id = ? ORDER BY created_at DESC LIMIT 20",
                    (conv_id,),
                )
            ]
            if len(messages) < 2:
                conn.close()
                return

            # 获取该用户+智能体已有的用户记忆
            existing = [
                dict(row)
                for row in conn.execute(
                    "SELECT * FROM user_memories WHERE user_id = ? AND agent_id = ?",
                    (user_id, agent_id),
                )
            ]
            existing_text = "\n".join(
                [f"[{m['category']}] {m['key']}: {m['value']}" for m in existing]
            ) if existing else "（暂无已记录的信息）"

            # 构建分析提示
            chat_text = "\n".join(
                [f"{'用户' if m['role']=='user' else '智能体'}: {m['content'][:200]}" for m in reversed(messages)]
            )

            analyze_prompt = f"""分析以下对话，提取关于**用户**的个人信息和偏好。

已有信息（请勿重复提取相同内容）：
{existing_text}

最近对话：
{chat_text}

⚠️ 提取规则（严格遵守）：
1. 只能提取用户**亲口说出**或**明确书面陈述**的信息
2. 禁止根据用户的用词风格、表情符号、提问方式等**推测**用户的性格、年龄、性别、外貌
3. 禁止编造用户的外貌长相、穿衣风格、爱好、生活习惯、饮食偏好、职业、家庭状况
4. 智能体对用户说的猜测性/恭维性话语（如"你看起来是个温柔的人"）不能作为提取依据
5. 只有用户说"我喜欢/我讨厌/我是/我住在/我在做..."这类明确陈述才能提取
6. 置信度 >= 0.8 才录入（即用户明确说出的信息），不要提取模糊暗示

请以 JSON 格式返回提取的新发现或更新的信息：
```json
[
  {{"category": "分类", "key": "信息标签", "value": "具体内容", "confidence": 0.0-1.0}}
]
```

分类可选：name, interests, personality_traits, life_context, preferences, goals, general
如果没有发现用户明确陈述的新信息，返回空数组 []。
只返回 JSON，不要其他内容。"""

            try:
                client = get_openai_client()
                response = client.chat.completions.create(
                    model=MODEL_NAME,
                    messages=[
                        {"role": "system", "content": "你是一个用户信息提取器。只返回 JSON 数组。"},
                        {"role": "user", "content": analyze_prompt},
                    ],
                    temperature=0.3,
                    max_tokens=800,
                )
                result_text = response.choices[0].message.content.strip()

                # 清理可能的 markdown 包裹
                if result_text.startswith("```"):
                    result_text = result_text.split("```")[1]
                    if result_text.startswith("json"):
                        result_text = result_text[4:]
                result_text = result_text.strip()

                items = json.loads(result_text)
                if not items:
                    conn.close()
                    return

                # 存入数据库（按智能体隔离）
                new_count = 0
                for item in items:
                    if item.get("confidence", 0) < 0.8:
                        continue
                    cat = item.get("category", "general")
                    key = item.get("key", "").strip()
                    val = item.get("value", "").strip()
                    if not key or not val:
                        continue

                    mem_id = str(uuid.uuid4())
                    # 使用 INSERT OR REPLACE 处理重复（同一智能体内）
                    conn.execute(
                        """INSERT OR REPLACE INTO user_memories (id, user_id, agent_id, category, key, value, source_conv_id, confidence, updated_at)
                           VALUES (?, ?, ?, ?, ?, ?, ?, ?, datetime('now'))""",
                        (mem_id, user_id, agent_id, cat, key, val, conv_id, item.get("confidence", 0.7)),
                    )
                    new_count += 1

                # 记录提取日志
                if new_count > 0:
                    summary = ", ".join([f"{it['key']}:{it['value'][:30]}" for it in items[:5]])
                    conn.execute(
                        "INSERT INTO extraction_log (id, user_id, conversation_id, extracted_count, extracted_summary) VALUES (?, ?, ?, ?, ?)",
                        (str(uuid.uuid4()), user_id, conv_id, new_count, summary),
                    )

                conn.commit()
            except (json.JSONDecodeError, Exception):
                pass  # 提取失败不影响主流程
        finally:
            conn.close()

    # 在新线程中执行，不阻塞回复
    t = threading.Thread(target=_do_extract, daemon=True)
    t.start()


# ==================== 用户认证工具 ====================
def _hash_password(password: str, salt: str = "") -> tuple:
    """密码哈希"""
    if not salt:
        salt = secrets.token_hex(16)
    h = hashlib.pbkdf2_hmac("sha256", password.encode("utf-8"), salt.encode("utf-8"), 100000)
    return h.hex(), salt


def _verify_password(password: str, salt: str, stored_hash: str) -> bool:
    """验证密码"""
    h, _ = _hash_password(password, salt)
    return h == stored_hash


def _get_current_user(conn, token: str) -> Optional[dict]:
    """根据 token 获取当前用户"""
    if not token:
        return None
    row = conn.execute("SELECT * FROM users WHERE token = ?", (token,)).fetchone()
    return dict(row) if row else None


def _require_user(authorization: Optional[str] = Header(None)) -> dict:
    """FastAPI 依赖：从 Authorization 头验证用户，要求登录"""
    if not authorization or not authorization.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="请先登录")
    token = authorization[7:]
    conn = get_db()
    try:
        user = _get_current_user(conn, token)
        if not user:
            raise HTTPException(status_code=401, detail="登录已过期，请重新登录")
        return user
    finally:
        conn.close()


def _require_admin(authorization: Optional[str] = Header(None)) -> dict:
    """FastAPI 依赖：验证用户是管理员"""
    user = _require_user(authorization)
    if not user.get("is_admin"):
        raise HTTPException(status_code=403, detail="仅管理员可执行此操作")
    return user


def _get_opt_user(authorization: Optional[str] = Header(None)) -> Optional[dict]:
    """FastAPI 依赖：可选的用户验证（未登录返回 None）"""
    if not authorization or not authorization.startswith("Bearer "):
        return None
    token = authorization[7:]
    conn = get_db()
    try:
        return _get_current_user(conn, token)
    finally:
        conn.close()


# ==================== Pydantic 模型 ====================
class AgentCreate(BaseModel):
    name: str
    personality: str = ""
    speaking_style: str = ""
    character_setting: str = ""
    avatar: str = "🤖"
    enable_search: bool = False


class AgentUpdate(BaseModel):
    name: Optional[str] = None
    personality: Optional[str] = None
    speaking_style: Optional[str] = None
    character_setting: Optional[str] = None
    avatar: Optional[str] = None
    enable_search: Optional[bool] = None


class ConversationCreate(BaseModel):
    agent_id: str
    title: str = "新对话"


class ConversationUpdate(BaseModel):
    title: Optional[str] = None


class ChatRequest(BaseModel):
    conversation_id: Optional[str] = None
    agent_id: str
    message: str
    enable_search: bool = True
    client_now: Optional[str] = None  # ISO 格式客户端本地时间
    location_city: Optional[str] = None  # 客户端当前城市（用于展示）
    latitude: Optional[float] = None  # GPS 纬度（用于天气查询）
    longitude: Optional[float] = None  # GPS 经度（用于天气查询）


class LocationUpdate(BaseModel):
    agent_id: str
    latitude: float
    longitude: float


class ChatResponse(BaseModel):
    conversation_id: str
    reply: str
    message_id: str


class RegisterRequest(BaseModel):
    username: str
    password: str
    nickname: str = ""


class LoginRequest(BaseModel):
    username: str
    password: str


class UserProfileUpdate(BaseModel):
    nickname: Optional[str] = None
    avatar_url: Optional[str] = None
    password: Optional[str] = None


class ApiConfig(BaseModel):
    api_key: str = ""
    api_base: str = "https://api.openai.com/v1"
    model: str = "gpt-3.5-turbo"


# ==================== FastAPI 应用 ====================
app = FastAPI(title="AI 聊天智能体", version="1.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# 全局客户端（在每次请求时更新）
openai_client: Optional[OpenAI] = None


def get_openai_client() -> OpenAI:
    global openai_client
    if openai_client is None:
        openai_client = OpenAI(api_key=API_KEY, base_url=API_BASE)
    return openai_client


# ==================== 初始化 ====================
@app.on_event("startup")
async def startup():
    init_db()
    # 启动提醒调度器后台线程（守护线程，随主进程退出）
    threading.Thread(target=_reminder_scheduler_loop, name="reminder-scheduler", daemon=True).start()


# ==================== 智能体 API ====================
@app.get("/api/agents")
def list_agents(user: dict = Depends(_require_user)):
    """获取智能体列表：系统内置 + 用户自建"""
    conn = get_db()
    try:
        agents = [
            dict(row)
            for row in conn.execute(
                "SELECT id, user_id, name, personality, speaking_style, character_setting, avatar, enable_search, created_at, updated_at FROM agents WHERE user_id = ? OR user_id = ? ORDER BY user_id = ? DESC, created_at DESC",
                (user["id"], SYSTEM_USER_ID, user["id"]),
            )
        ]
        return {"agents": agents}
    finally:
        conn.close()


@app.get("/api/agents/{agent_id}")
def get_agent(agent_id: str, user: dict = Depends(_require_user)):
    """获取单个智能体"""
    conn = get_db()
    try:
        row = conn.execute(
            "SELECT id, user_id, name, personality, speaking_style, character_setting, avatar, enable_search, created_at, updated_at FROM agents WHERE id = ? AND (user_id = ? OR user_id = ?)",
            (agent_id, user["id"], SYSTEM_USER_ID),
        ).fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="智能体不存在")
        return {"agent": dict(row)}
    finally:
        conn.close()


@app.get("/api/agents/{agent_id}/sleep-window")
def get_sleep_window(agent_id: str, user: dict = Depends(_require_user)):
    """获取智能体的睡眠窗口（根据用户画像动态计算）"""
    conn = get_db()
    try:
        # 验证智能体属于用户或是系统内置
        row = conn.execute("SELECT id FROM agents WHERE id = ? AND (user_id = ? OR user_id = ?)", (agent_id, user["id"], SYSTEM_USER_ID)).fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="智能体不存在")
        return _get_sleep_window(conn, agent_id, user["id"])
    finally:
        conn.close()


@app.post("/api/agents")
def create_agent(data: AgentCreate, user: dict = Depends(_require_user)):
    """创建新智能体"""
    agent_id = str(uuid.uuid4())
    system_prompt = _build_system_prompt(data.name, data.personality, data.speaking_style, data.character_setting)
    conn = get_db()
    try:
        conn.execute(
            """INSERT INTO agents (id, user_id, name, personality, speaking_style, character_setting, avatar, system_prompt, enable_search)
               VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)""",
            (agent_id, user["id"], data.name, data.personality, data.speaking_style, data.character_setting, data.avatar, system_prompt, int(data.enable_search)),
        )
        conn.commit()
        row = conn.execute("SELECT * FROM agents WHERE id = ?", (agent_id,)).fetchone()
        return {"agent": dict(row)}
    finally:
        conn.close()


@app.put("/api/agents/{agent_id}")
def update_agent(agent_id: str, data: AgentUpdate, user: dict = Depends(_require_user)):
    """更新智能体（只能更新自己的）"""
    conn = get_db()
    try:
        existing = conn.execute("SELECT * FROM agents WHERE id = ? AND user_id = ?", (agent_id, user["id"])).fetchone()
        if not existing:
            # 检查是否是系统智能体
            sys_agent = conn.execute("SELECT * FROM agents WHERE id = ? AND user_id = ?", (agent_id, SYSTEM_USER_ID)).fetchone()
            if sys_agent:
                raise HTTPException(status_code=403, detail="不能修改系统内置智能体")
            raise HTTPException(status_code=404, detail="智能体不存在")

        updates = {}
        if data.name is not None:
            updates["name"] = data.name
        if data.personality is not None:
            updates["personality"] = data.personality
        if data.speaking_style is not None:
            updates["speaking_style"] = data.speaking_style
        if data.character_setting is not None:
            updates["character_setting"] = data.character_setting
        if data.avatar is not None:
            updates["avatar"] = data.avatar
        if data.enable_search is not None:
            updates["enable_search"] = int(data.enable_search)

        if updates:
            # 重新构建 system prompt
            name = updates.get("name", existing["name"])
            personality = updates.get("personality", existing["personality"])
            style = updates.get("speaking_style", existing["speaking_style"])
            char_setting = updates.get("character_setting", existing["character_setting"])
            updates["system_prompt"] = _build_system_prompt(name, personality, style, char_setting)
            updates["updated_at"] = datetime.now().astimezone().isoformat()

            set_clause = ", ".join(f"{k} = ?" for k in updates)
            values = list(updates.values()) + [agent_id]
            conn.execute(f"UPDATE agents SET {set_clause} WHERE id = ?", values)
            conn.commit()

        row = conn.execute("SELECT * FROM agents WHERE id = ?", (agent_id,)).fetchone()
        return {"agent": dict(row)}
    finally:
        conn.close()


@app.delete("/api/agents/{agent_id}")
def delete_agent(agent_id: str, user: dict = Depends(_require_user)):
    """删除智能体（只能删除自己的；同时清理知识库文件和向量数据）"""
    conn = get_db()
    try:
        existing = conn.execute("SELECT * FROM agents WHERE id = ? AND user_id = ?", (agent_id, user["id"])).fetchone()
        if not existing:
            sys_agent = conn.execute("SELECT * FROM agents WHERE id = ? AND user_id = ?", (agent_id, SYSTEM_USER_ID)).fetchone()
            if sys_agent:
                raise HTTPException(status_code=403, detail="不能删除系统内置智能体")
            raise HTTPException(status_code=404, detail="智能体不存在")
        # 清理 ChromaDB 向量数据
        try:
            client = _get_chroma_client()
            safe_name = f"kb_{agent_id.replace('-', '_')}"
            try:
                client.delete_collection(name=safe_name)
                logger.info(f"已删除智能体 {agent_id} 的向量集合")
            except Exception:
                pass  # 集合可能不存在
        except Exception:
            pass

        # 清理本地知识库文件
        kb_dir = BASE_DIR / "data" / "knowledge" / agent_id
        if kb_dir.exists():
            shutil.rmtree(kb_dir, ignore_errors=True)

        conn.execute("DELETE FROM agents WHERE id = ?", (agent_id,))
        conn.commit()
        return {"success": True}
    finally:
        conn.close()


# ==================== 对话 API ====================
@app.get("/api/conversations")
def list_conversations(agent_id: Optional[str] = None, user: dict = Depends(_require_user)):
    """获取对话列表（包含最后一条消息内容）"""
    conn = get_db()
    try:
        base_sql = """SELECT c.*, a.name as agent_name, a.avatar as agent_avatar,
                          (SELECT content FROM messages WHERE conversation_id = c.id ORDER BY created_at DESC LIMIT 1) as last_message
                   FROM conversations c
                   JOIN agents a ON c.agent_id = a.id
                   WHERE c.user_id = ?"""
        params = [user["id"]]
        if agent_id:
            base_sql += " AND c.agent_id = ?"
            params.append(agent_id)
        base_sql += " ORDER BY c.updated_at DESC"
        rows = conn.execute(base_sql, params)
        return {"conversations": [dict(row) for row in rows]}
    finally:
        conn.close()


@app.get("/api/conversations/{conversation_id}")
def get_conversation(conversation_id: str, user: dict = Depends(_require_user)):
    """获取对话详情"""
    conn = get_db()
    try:
        conv = conn.execute(
            "SELECT * FROM conversations WHERE id = ? AND user_id = ?", (conversation_id, user["id"])
        ).fetchone()
        if not conv:
            raise HTTPException(status_code=404, detail="对话不存在")

        messages = [
            dict(row)
            for row in conn.execute(
                "SELECT * FROM messages WHERE conversation_id = ? ORDER BY created_at ASC",
                (conversation_id,),
            )
        ]
        return {"conversation": dict(conv), "messages": messages}
    finally:
        conn.close()


@app.post("/api/conversations")
def create_conversation(data: ConversationCreate, user: dict = Depends(_require_user)):
    """创建新对话"""
    conv_id = str(uuid.uuid4())
    conn = get_db()
    try:
        conn.execute(
            "INSERT INTO conversations (id, agent_id, user_id, title) VALUES (?, ?, ?, ?)",
            (conv_id, data.agent_id, user["id"], data.title),
        )
        conn.commit()
        row = conn.execute("SELECT * FROM conversations WHERE id = ?", (conv_id,)).fetchone()
        return {"conversation": dict(row)}
    finally:
        conn.close()


@app.delete("/api/conversations/{conversation_id}")
def delete_conversation(conversation_id: str, user: dict = Depends(_require_user)):
    """删除对话"""
    conn = get_db()
    try:
        conn.execute("DELETE FROM conversations WHERE id = ? AND user_id = ?", (conversation_id, user["id"]))
        conn.commit()
        return {"success": True}
    finally:
        conn.close()




# ==================== 批量消息 API（用于快速插入预定义消息，无需 AI 生成） ====================
class BatchMessagesRequest(BaseModel):
    messages: list = []  # [{"role": "assistant", "content": "..."}, ...]


@app.post("/api/conversations/{conv_id}/messages/batch")
def batch_create_messages(conv_id: str, data: BatchMessagesRequest, user: dict = Depends(_require_user)):
    """批量插入预定义消息到对话（跳过 AI 生成，直接写入）"""
    conn = get_db()
    try:
        # 校验对话属于当前用户
        conv = conn.execute(
            "SELECT * FROM conversations WHERE id = ? AND user_id = ?", (conv_id, user["id"])
        ).fetchone()
        if not conv:
            raise HTTPException(status_code=404, detail="对话不存在")

        now = datetime.now().astimezone().isoformat()
        inserted = []
        for msg in data.messages:
            msg_id = str(uuid.uuid4())
            conn.execute(
                "INSERT INTO messages (id, conversation_id, role, content, created_at) VALUES (?, ?, ?, ?, ?)",
                (msg_id, conv_id, msg.get("role", "assistant"), msg["content"], now),
            )
            inserted.append({"id": msg_id, "role": msg["role"], "content": msg["content"], "created_at": now})

        # 更新对话的 updated_at
        conn.execute("UPDATE conversations SET updated_at = ? WHERE id = ?", (now, conv_id))
        conn.commit()
        return {"messages": inserted}
    finally:
        conn.close()


# ==================== 聊天 API ====================
@app.post("/api/chat")
async def chat(data: ChatRequest, user: dict = Depends(_require_user)):
    """发送消息并获取 AI 回复"""
    if not API_KEY:
        raise HTTPException(status_code=500, detail="请先配置 API Key")

    conn = get_db()
    try:
        # 智能体必须属于用户或是系统内置的
        agent = conn.execute(
            "SELECT * FROM agents WHERE id = ? AND (user_id = ? OR user_id = ?)", (data.agent_id, user["id"], SYSTEM_USER_ID)
        ).fetchone()
        if not agent:
            raise HTTPException(status_code=404, detail="智能体不存在")

        # 获取或创建对话
        conversation_id = data.conversation_id
        if not conversation_id:
            conversation_id = str(uuid.uuid4())
            conn.execute(
                "INSERT INTO conversations (id, agent_id, user_id, title) VALUES (?, ?, ?, ?)",
                (conversation_id, data.agent_id, user["id"], data.message[:30] + ("..." if len(data.message) > 30 else "")),
            )
        else:
            # 验证对话属于当前用户
            conv = conn.execute("SELECT id FROM conversations WHERE id = ? AND user_id = ?", (conversation_id, user["id"])).fetchone()
            if not conv:
                raise HTTPException(status_code=404, detail="对话不存在")
            # 更新对话时间（使用客户端时间，与时区保持一致）
            conn.execute(
                "UPDATE conversations SET updated_at = ? WHERE id = ? AND user_id = ?",
                (data.client_now or datetime.now().isoformat(), conversation_id, user["id"]),
            )

        # 保存用户消息（使用客户端带时区的时间）
        user_msg_id = str(uuid.uuid4())
        conn.execute(
            "INSERT INTO messages (id, conversation_id, role, content, created_at) VALUES (?, ?, ?, ?, ?)",
            (user_msg_id, conversation_id, "user", data.message, data.client_now or datetime.now().isoformat()),
        )
        conn.commit()

        # 获取最近30条历史消息（避免上下文过长导致重复或偏离）
        history = [
            dict(row)
            for row in conn.execute(
                "SELECT role, content FROM (SELECT role, content, created_at FROM messages WHERE conversation_id = ? ORDER BY created_at DESC LIMIT 30) ORDER BY created_at ASC",
                (conversation_id,),
            )
        ]

        # 构建注入用户画像的 system prompt（仅本智能体的记忆）
        user_profile_text = _get_user_profile_text(conn, data.agent_id, user["id"])

        # 获取时间上下文（优先使用客户端本地时间，避免服务器时区偏差）
        time_context = _get_time_context(data.client_now)

        # 根据智能体的 enable_search 设置决定是否联网搜索
        agent_enable_search = bool(agent["enable_search"])
        search_context = ""
        weather_context = ""
        if agent_enable_search and data.message.strip():
            search_context = _web_search(data.message)
            # 获取天气信息：优先用客户端城市，其次数据库记录
            location = data.location_city or _get_user_location(conn, data.agent_id, user["id"]) or "北京"
            weather_context = _fetch_weather(location, data.latitude, data.longitude)
            # 微博热搜（实时话题素材）
            weibo_hot = _fetch_weibo_hot(10)
            if weibo_hot:
                search_context += "\n\n" + weibo_hot

        # RAG 知识库检索
        knowledge_context = ""
        try:
            relevant_chunks = query_knowledge_base(data.agent_id, data.message, user_id=user["id"])
            if relevant_chunks:
                knowledge_context = "\n\n---\n".join(
                    f"[资料片段 {i+1}]\n{chunk}"
                    for i, chunk in enumerate(relevant_chunks)
                )
                if knowledge_context:
                    knowledge_context = f"以下是你记忆中的相关资料（共 {len(relevant_chunks)} 条）：\n\n{knowledge_context}"
        except Exception as e:
            logger.warning(f"知识库检索跳过: {e}")

        enhanced_prompt = _build_system_prompt(
            agent["name"],
            agent["personality"],
            agent["speaking_style"],
            agent["character_setting"],
            user_profile_text,
            time_context,
            search_context,
            knowledge_context,
            weather_context,
        )

        # 构建消息列表
        messages = [{"role": "system", "content": enhanced_prompt}]
        for msg in history:
            messages.append({"role": msg["role"], "content": msg["content"]})

        # 调用 AI（知识库模式下多给一些 tokens 用于引用资料）
        max_tokens = 1000 if knowledge_context else 600
        client = get_openai_client()
        try:
            response = client.chat.completions.create(
                model=MODEL_NAME,
                messages=messages,
                temperature=0.7,
                max_tokens=max_tokens,
                timeout=120.0,  # 🔥 120秒超时，防止 API 挂死
            )
            reply = response.choices[0].message.content
        except Exception as e:
            reply = f"抱歉，我暂时无法回复。(错误: {str(e)})"

        # 保存 AI 回复（先解析并剥离提醒标签）
        clean_reply, reminders = _parse_reminders(reply)
        ai_msg_id = str(uuid.uuid4())
        conn.execute(
            "INSERT INTO messages (id, conversation_id, role, content, created_at) VALUES (?, ?, ?, ?, ?)",
            (ai_msg_id, conversation_id, "assistant", clean_reply, data.client_now or datetime.now().isoformat()),
        )
        conn.commit()

        # 将提取到的提醒任务写入 reminders 表
        for r in reminders:
            try:
                # 验证时间格式
                remind_at_dt = datetime.fromisoformat(r["remind_at"])
                if remind_at_dt > datetime.now():
                    reminder_id = str(uuid.uuid4())
                    conn.execute(
                        "INSERT INTO reminders (id, agent_id, user_id, conversation_id, remind_at, content) VALUES (?, ?, ?, ?, ?, ?)",
                        (reminder_id, data.agent_id, user["id"], conversation_id, r["remind_at"], r["content"]),
                    )
                    conn.commit()
                    logger.info(f"提醒已创建: {r['remind_at']} — {r['content']}")
                else:
                    logger.info(f"提醒时间已过期，跳过: {r['remind_at']}")
            except Exception as e:
                logger.warning(f"提醒解析失败: {e}")

        # 异步提取用户偏好（后台线程，不阻塞回复，按智能体隔离）
        _extract_user_preferences_async(conversation_id, data.agent_id, user["id"])

        return {
            "conversation_id": conversation_id,
            "reply": clean_reply,
            "message_id": ai_msg_id,
        }
    finally:
        conn.close()


# ==================== 智能体主动消息 API ====================
class InitiativeRequest(BaseModel):
    agent_id: str
    conversation_id: Optional[str] = None
    enable_search: bool = True
    is_intro: bool = False
    client_now: Optional[str] = None  # ISO 格式客户端本地时间
    location_city: Optional[str] = None  # 客户端当前城市（用于展示）
    latitude: Optional[float] = None  # GPS 纬度（用于天气查询）
    longitude: Optional[float] = None  # GPS 经度（用于天气查询）


@app.post("/api/agents/initiative-message")
def get_initiative_message(data: InitiativeRequest, user: dict = Depends(_require_user)):
    """让智能体主动发起消息（1~2条）"""
    if not API_KEY:
        raise HTTPException(status_code=500, detail="请先配置 API Key")

    conn = get_db()
    try:
        agent = conn.execute("SELECT * FROM agents WHERE id = ? AND (user_id = ? OR user_id = ?)", (data.agent_id, user["id"], SYSTEM_USER_ID)).fetchone()
        if not agent:
            raise HTTPException(status_code=404, detail="智能体不存在")

        # 获取最近几条消息作为上下文（反过来再 reverse 得到时间升序）
        history = []
        if data.conversation_id:
            # 🔒 安全校验：确保 conversation 属于当前用户
            conv = conn.execute(
                "SELECT id FROM conversations WHERE id = ? AND user_id = ?",
                (data.conversation_id, user["id"]),
            ).fetchone()
            if not conv:
                conn.close()
                raise HTTPException(status_code=404, detail="对话不存在")
            rows = conn.execute(
                "SELECT role, content FROM messages WHERE conversation_id = ? ORDER BY created_at DESC LIMIT 6",
                (data.conversation_id,),
            ).fetchall()
            history = [dict(r) for r in reversed(rows)]

        # 构建上下文
        user_profile = _get_user_profile_text(conn, data.agent_id, user["id"])
        time_context = _get_time_context(data.client_now)

        # 🔥 联网搜索：根据智能体设置决定是否联网
        agent_enable_search = bool(agent["enable_search"])
        search_context = ""
        weather_context = ""
        if agent_enable_search:
            # 优先用最近用户消息作为搜索词，否则用通用时政关键词
            search_query = ""
            for h in reversed(history):
                if h["role"] == "user":
                    search_query = h["content"][:100]
                    break
            if not search_query:
                search_query = "今日热点新闻"
            search_context = _web_search(search_query)
            # 获取天气：优先用客户端城市，其次数据库记录
            location = data.location_city or _get_user_location(conn, data.agent_id, user["id"]) or "北京"
            weather_context = _fetch_weather(location, data.latitude, data.longitude)
            # 微博热搜（实时话题素材）
            weibo_hot = _fetch_weibo_hot(10)
            if weibo_hot:
                search_context += "\n\n" + weibo_hot

        initiative_prompt = _build_initiative_prompt(
            agent["name"], agent["personality"], agent["speaking_style"],
            agent["character_setting"], user_profile, time_context,
            agent_enable_search, data.is_intro, search_context, weather_context,
        )

        messages = [{"role": "system", "content": initiative_prompt}]
        for msg in history:
            messages.append({"role": msg["role"], "content": msg["content"]})

        # 注入一条 user 消息作为触发
        messages.append({"role": "user", "content": "（现在请你主动发起对话，发1~2条消息。直接输出消息内容，每条独立成段，不要加任何前缀或说明。）"})

        client = get_openai_client()
        try:
            response = client.chat.completions.create(
                model=MODEL_NAME,
                messages=messages,
                temperature=0.85,
                max_tokens=300,
                timeout=60.0,  # 🔥 60秒超时，防止挂死
            )
            reply = response.choices[0].message.content
        except Exception as e:
            return {"message_id": "", "message": "", "error": str(e)}

        # 分隔多条消息（按空行拆分）
        raw_msgs = [m.strip() for m in reply.split("\n\n") if m.strip()]
        if not raw_msgs:
            raw_msgs = [reply.strip()]

        # 最多取 2 条
        raw_msgs = raw_msgs[:2]

        # 保存到数据库
        saved = []
        now_iso = data.client_now or datetime.now().astimezone().isoformat()
        for text in raw_msgs:
            ai_msg_id = str(uuid.uuid4())
            if data.conversation_id:
                conn.execute(
                    "INSERT INTO messages (id, conversation_id, role, content, created_at) VALUES (?, ?, ?, ?, ?)",
                    (ai_msg_id, data.conversation_id, "assistant", text, now_iso),
                )
            saved.append({"id": ai_msg_id, "role": "assistant", "content": text})

        if data.conversation_id:
            conn.execute(
                "UPDATE conversations SET updated_at = ? WHERE id = ?",
                (now_iso, data.conversation_id),
            )
            conn.commit()

        return {
            "conversation_id": data.conversation_id,
            "messages": saved,
        }
    finally:
        conn.close()


# ==================== 配置 API ====================
@app.post("/api/config")
def save_config(config: ApiConfig, _admin: dict = Depends(_require_admin)):
    """保存 API 配置到 .env 文件"""
    env_path = BASE_DIR / ".env"
    lines = []
    if os.path.exists(env_path):
        with open(env_path, "r", encoding="utf-8") as f:
            lines = f.readlines()

    env_vars = {}
    # 只有提供了才更新
    if config.api_key.strip():
        env_vars["OPENAI_API_KEY"] = config.api_key.strip()
    if config.api_base.strip():
        env_vars["OPENAI_API_BASE"] = config.api_base.strip()
    if config.model.strip():
        env_vars["OPENAI_MODEL"] = config.model.strip()

    for key, value in env_vars.items():
        found = False
        for i, line in enumerate(lines):
            if line.startswith(f"{key}="):
                lines[i] = f"{key}={value}\n"
                found = True
                break
        if not found:
            lines.append(f"{key}={value}\n")

    with open(env_path, "w", encoding="utf-8") as f:
        f.writelines(lines)

    # 重新加载环境变量
    load_dotenv(env_path, override=True)

    # 更新全局变量
    global API_KEY, API_BASE, MODEL_NAME, openai_client
    API_KEY = os.getenv("OPENAI_API_KEY", API_KEY)
    API_BASE = os.getenv("OPENAI_API_BASE", API_BASE)
    MODEL_NAME = os.getenv("OPENAI_MODEL", MODEL_NAME)
    openai_client = OpenAI(api_key=API_KEY, base_url=API_BASE)

    return {"success": True, "message": "配置已保存"}


@app.get("/api/config")
def get_config(user: dict = Depends(_require_user)):
    """获取当前 API 配置（隐藏 key 部分）"""
    masked_key = ""
    if API_KEY:
        masked_key = API_KEY[:8] + "****" + API_KEY[-4:] if len(API_KEY) > 12 else "****"
    return {
        "api_key": masked_key,
        "api_base": API_BASE,
        "model": MODEL_NAME,
        "has_key": bool(API_KEY),
    }


# ==================== 搜索 API ====================
class SearchRequest(BaseModel):
    query: str


@app.post("/api/search")
def search_web(data: SearchRequest, user: dict = Depends(_require_user)):
    """联网搜索"""
    if not data.query.strip():
        raise HTTPException(status_code=400, detail="请输入搜索内容")
    results = _web_search(data.query)
    return {"success": True, "results": results, "query": data.query}


# ==================== 用户认证 API ====================
@app.post("/api/auth/register")
def register_user(data: RegisterRequest):
    """用户注册"""
    username = data.username.strip()
    password = data.password.strip()
    nickname = data.nickname.strip() or username

    if not username or not password:
        raise HTTPException(status_code=400, detail="用户名和密码不能为空")
    if len(username) < 2 or len(username) > 20:
        raise HTTPException(status_code=400, detail="用户名长度应为 2-20 个字符")
    if len(password) < 4:
        raise HTTPException(status_code=400, detail="密码至少需要 4 个字符")

    conn = get_db()
    try:
        existing = conn.execute("SELECT id FROM users WHERE username = ?", (username,)).fetchone()
        if existing:
            raise HTTPException(status_code=400, detail="用户名已被使用")

        user_id = str(uuid.uuid4())
        h, salt = _hash_password(password)
        token = secrets.token_hex(32)
        is_admin = 1 if ADMIN_USERNAME and username == ADMIN_USERNAME else 0

        conn.execute(
            """INSERT INTO users (id, username, password_hash, salt, nickname, token, is_admin, created_at, updated_at)
               VALUES (?, ?, ?, ?, ?, ?, ?, datetime('now'), datetime('now'))""",
            (user_id, username, h, salt, nickname, token, is_admin),
        )
        conn.commit()

        return {
            "success": True,
            "user": {
                "id": user_id,
                "username": username,
                "nickname": nickname,
                "avatar_url": "",
                "token": token,
                "is_admin": bool(is_admin),
            },
        }
    finally:
        conn.close()


@app.post("/api/auth/login")
def login_user(data: LoginRequest):
    """用户登录"""
    username = data.username.strip()
    password = data.password.strip()

    if not username or not password:
        raise HTTPException(status_code=400, detail="请输入用户名和密码")

    conn = get_db()
    try:
        row = conn.execute("SELECT * FROM users WHERE username = ?", (username,)).fetchone()
        if not row:
            raise HTTPException(status_code=401, detail="用户名或密码错误")
        user = dict(row)

        if not _verify_password(password, user["salt"], user["password_hash"]):
            raise HTTPException(status_code=401, detail="用户名或密码错误")

        # 每次登录生成新 token
        token = secrets.token_hex(32)
        conn.execute(
            "UPDATE users SET token = ?, updated_at = datetime('now') WHERE id = ?",
            (token, user["id"]),
        )
        conn.commit()

        return {
            "success": True,
            "user": {
                "id": user["id"],
                "username": user["username"],
                "nickname": user["nickname"],
                "avatar_url": user["avatar_url"],
                "token": token,
                "is_admin": bool(user.get("is_admin", 0)),
            },
        }
    finally:
        conn.close()


@app.post("/api/auth/validate")
def validate_token(request: Request):
    """验证 token 是否有效"""
    token = request.headers.get("authorization", "").replace("Bearer ", "")
    if not token:
        return {"valid": False}
    conn = get_db()
    try:
        user = _get_current_user(conn, token)
        if user:
            return {
                "valid": True,
                "user": {
                    "id": user["id"],
                    "username": user["username"],
                    "nickname": user["nickname"],
                    "avatar_url": user["avatar_url"],
                    "is_admin": bool(user.get("is_admin", 0)),
                },
            }
        return {"valid": False}
    finally:
        conn.close()


@app.post("/api/user/update-profile")
def update_user_profile(data: UserProfileUpdate, user: dict = Depends(_require_user)):
    """更新用户个人资料（昵称/头像/密码）"""
    conn = get_db()
    try:

        updates = []
        values = []

        if data.nickname is not None and data.nickname.strip():
            updates.append("nickname = ?")
            values.append(data.nickname.strip())

        if data.avatar_url is not None:
            updates.append("avatar_url = ?")
            values.append(data.avatar_url)

        if data.password is not None and data.password.strip():
            h, salt = _hash_password(data.password.strip())
            updates.append("password_hash = ?")
            values.append(h)
            updates.append("salt = ?")
            values.append(salt)

        if updates:
            updates.append("updated_at = datetime('now')")
            values.append(user["id"])
            conn.execute(f"UPDATE users SET {', '.join(updates)} WHERE id = ?", values)
            conn.commit()

        # 返回更新后的用户信息
        row = conn.execute("SELECT * FROM users WHERE id = ?", (user["id"],)).fetchone()
        u = dict(row)
        return {
            "success": True,
            "user": {
                "id": u["id"],
                "username": u["username"],
                "nickname": u["nickname"],
                "avatar_url": u["avatar_url"],
            },
        }
    finally:
        conn.close()


# ==================== 用户画像 API ====================
class MemoryUpdate(BaseModel):
    category: str = "general"
    key: str
    value: str
    confidence: float = 0.7
    agent_id: str = ""


@app.get("/api/user/profile")
def get_user_profile(agent_id: str = "", user: dict = Depends(_require_user)):
    """获取用户画像（可按智能体过滤）"""
    conn = get_db()
    try:
        params = [user["id"]]
        where_clause = "WHERE user_id = ?"
        if agent_id:
            where_clause += " AND agent_id = ?"
            params.append(agent_id)
        memories = [
            dict(row)
            for row in conn.execute(
                f"SELECT * FROM user_memories {where_clause} ORDER BY category, key",
                params,
            )
        ]

        # 获取提取日志（🔒 仅当前用户的日志）
        logs = [
            dict(row)
            for row in conn.execute(
                "SELECT * FROM extraction_log WHERE user_id = ? ORDER BY created_at DESC LIMIT 20",
                (user["id"],),
            )
        ]

        profile_text = _get_user_profile_text(conn, agent_id, user["id"])

        return {
            "memories": memories,
            "profile_text": profile_text,
            "extraction_logs": logs,
            "total_memories": len(memories),
        }
    finally:
        conn.close()


@app.post("/api/user/memory")
def add_user_memory(data: MemoryUpdate, user: dict = Depends(_require_user)):
    """手动添加/更新一条用户记忆（按用户+智能体隔离）"""
    conn = get_db()
    try:
        mem_id = str(uuid.uuid4())
        conn.execute(
            """INSERT OR REPLACE INTO user_memories (id, user_id, agent_id, category, key, value, confidence, updated_at)
               VALUES (?, ?, ?, ?, ?, ?, ?, datetime('now'))""",
            (mem_id, user["id"], data.agent_id, data.category, data.key, data.value, data.confidence),
        )
        conn.commit()
        return {"success": True, "memory_id": mem_id}
    finally:
        conn.close()


@app.delete("/api/user/memory/{memory_id}")
def delete_user_memory(memory_id: str, user: dict = Depends(_require_user)):
    """删除一条用户记忆"""
    conn = get_db()
    try:
        conn.execute("DELETE FROM user_memories WHERE id = ? AND user_id = ?", (memory_id, user["id"]))
        conn.commit()
        return {"success": True}
    finally:
        conn.close()


@app.post("/api/user/extract-now")
def extract_user_preferences_now(agent_id: str = "", user: dict = Depends(_require_user)):
    """手动触发用户偏好提取（按用户+智能体隔离）"""
    conn = get_db()
    try:
        # 获取该用户+智能体的对话消息
        if agent_id:
            all_messages = [
                dict(row)
                for row in conn.execute(
                    """SELECT m.role, m.content, m.conversation_id FROM messages m
                       JOIN conversations c ON m.conversation_id = c.id
                       WHERE c.agent_id = ? AND c.user_id = ? ORDER BY m.created_at DESC LIMIT 100""",
                    (agent_id, user["id"]),
                )
            ]
        else:
            all_messages = [
                dict(row)
                for row in conn.execute(
                    """SELECT m.role, m.content, m.conversation_id FROM messages m
                       JOIN conversations c ON m.conversation_id = c.id
                       WHERE c.user_id = ? ORDER BY m.created_at DESC LIMIT 100""",
                    (user["id"],),
                )
            ]

        if not all_messages:
            return {"success": False, "message": "没有对话历史"}

        existing = [
            dict(row)
            for row in conn.execute(
                "SELECT * FROM user_memories WHERE user_id = ? AND agent_id = ?",
                (user["id"], agent_id),
            )
        ]
        existing_text = "\n".join(
            [f"[{m['category']}] {m['key']}: {m['value']}" for m in existing]
        ) if existing else "（暂无已记录的信息）"

        chat_text = "\n".join([
            f"{'用户' if m['role']=='user' else 'AI'}: {m['content'][:200]}"
            for m in reversed(all_messages)
        ])

        analyze_prompt = f"""全面分析以下所有对话历史，提取关于**用户**的个人信息和偏好。

已有信息（请勿重复提取相同内容）：
{existing_text}

所有对话：
{chat_text}

⚠️ 提取规则（严格遵守）：
1. 只能提取用户**亲口说出**或**明确书面陈述**的信息
2. 禁止根据用户的用词风格、表情符号、提问方式等**推测**用户的性格、年龄、性别、外貌
3. 禁止编造用户的外貌长相、穿衣风格、爱好、生活习惯、饮食偏好、职业、家庭状况
4. 智能体对用户说的猜测性/恭维性话语（如"你看起来是个温柔的人"）不能作为提取依据
5. 只有用户说"我喜欢/我讨厌/我是/我住在/我在做..."这类明确陈述才能提取
6. 置信度 >= 0.8 才录入（即用户明确说出的信息），不要提取模糊暗示

请以 JSON 格式返回提取的新发现或更新的信息：
```json
[
  {{"category": "分类", "key": "信息标签", "value": "具体内容", "confidence": 0.0-1.0}}
]
```

分类可选：name, interests, personality_traits, life_context, preferences, goals, general
如果没有发现用户明确陈述的新信息，返回空数组 []。
只返回 JSON，不要其他内容。"""

        client = get_openai_client()
        response = client.chat.completions.create(
            model=MODEL_NAME,
            messages=[
                {"role": "system", "content": "你是用户信息提取器。只返回 JSON 数组。"},
                {"role": "user", "content": analyze_prompt},
            ],
            temperature=0.3,
            max_tokens=800,
        )
        result_text = response.choices[0].message.content.strip()

        if result_text.startswith("```"):
            result_text = result_text.split("```")[1]
            if result_text.startswith("json"):
                result_text = result_text[4:]
        result_text = result_text.strip()

        items = json.loads(result_text)
        new_count = 0
        for item in items:
            if item.get("confidence", 0) < 0.8:
                continue
            cat = item.get("category", "general")
            key = item.get("key", "").strip()
            val = item.get("value", "").strip()
            if not key or not val:
                continue

            mem_id = str(uuid.uuid4())
            conn.execute(
                """INSERT OR REPLACE INTO user_memories (id, user_id, agent_id, category, key, value, confidence, updated_at)
                   VALUES (?, ?, ?, ?, ?, ?, ?, datetime('now'))""",
                (mem_id, user["id"], agent_id, cat, key, val, item.get("confidence", 0.7)),
            )
            new_count += 1

        if new_count > 0:
            summary = ", ".join([f"{it['key']}:{it['value'][:30]}" for it in items[:5]])
            conn.execute(
                "INSERT INTO extraction_log (id, user_id, conversation_id, extracted_count, extracted_summary) VALUES (?, ?, ?, ?, ?)",
                (str(uuid.uuid4()), user["id"], "manual", new_count, summary),
            )
            conn.commit()

        return {
            "success": True,
            "new_memories": new_count,
            "summary": summary if new_count > 0 else "没有新发现",
        }
    except json.JSONDecodeError:
        return {"success": False, "message": f"解析失败: {result_text[:200]}"}
    except Exception as e:
        return {"success": False, "message": str(e)}
    finally:
        conn.close()


@app.post("/api/user/location")
def update_user_location(data: LocationUpdate, user: dict = Depends(_require_user)):
    """接收浏览器地理位置（经纬度），逆地理编码获取城市名，存入用户画像，用于天气提醒"""
    conn = get_db()
    try:
        # 1. 逆地理编码，获取城市名（使用免费的 OpenStreetMap Nominatim）
        city = ""
        try:
            import urllib.request, json as _json
            # zoom=16 精度约 100m~500m，比 zoom=10（~150km）精准得多
            nom_url = (
                f"https://nominatim.openstreetmap.org/reverse"
                f"?format=json&lat={data.latitude}&lon={data.longitude}&zoom=16&accept-language=zh"
            )
            req = urllib.request.Request(nom_url, headers={
                "User-Agent": "AI-Chat-Agent/1.0 (weather reminder only)"
            })
            with urllib.request.urlopen(req, timeout=8) as resp:
                addr = _json.loads(resp.read().decode("utf-8"))
                address_parts = addr.get("address", {})
                logger.info(f"Nominatim 返回地址: {address_parts}")
                # 中文地址：state=省、city=市/区、county=县/区、town=镇/街道
                raw_city = address_parts.get("city", "")
                state = address_parts.get("state", "")
                county = address_parts.get("county", "")
                town = address_parts.get("town", "")

                # 智能选城：优先用 city（通常为地级市），为空则逐级降级
                # 中国地址中 county 有时就是区名（如"南山区"），此时用 city 更合适
                if raw_city:
                    # 如果 city 以"市"结尾，说明是完整城市名
                    # 如果 city 不含"市"（可能是区名），先看 county 是否为市
                    if not raw_city.endswith("市") and county and county.endswith("市"):
                        city = county
                    else:
                        city = raw_city
                elif county:
                    city = county
                elif town:
                    city = town
                elif state:
                    city = state
                else:
                    city = addr.get("display_name", "")[:30]

                # 清洗城市名：去除行政后缀，wttr.in 对纯城市名识别更准
                for suffix in ("市", "省", "自治区", "特别行政区"):
                    if city.endswith(suffix) and len(city) > len(suffix):
                        city = city[:-len(suffix)]
                        break
        except Exception as e:
            logger.warning(f"逆地理编码失败: {e}")
            # Nominatim 失败 → 尝试 IP 地理定位兜底
            try:
                ip_url = "http://ip-api.com/json/?lang=zh-CN"
                ip_req = urllib.request.Request(ip_url, headers={
                    "User-Agent": "AI-Chat-Agent/1.0"
                })
                with urllib.request.urlopen(ip_req, timeout=5) as ip_resp:
                    ip_data = _json.loads(ip_resp.read().decode("utf-8"))
                    city = ip_data.get("city") or ip_data.get("regionName") or ""
                    if not city:
                        city = f"{data.latitude:.2f},{data.longitude:.2f}"
            except Exception as e2:
                logger.warning(f"IP 地理定位也失败: {e2}")
                city = f"{data.latitude:.2f},{data.longitude:.2f}"

        # 2. 存储到 user_memories（key 固定为"地理位置"，方便统一查询）
        mem_id = str(uuid.uuid4())
        conn.execute(
            """INSERT OR REPLACE INTO user_memories (id, user_id, agent_id, category, key, value, confidence, updated_at)
               VALUES (?, ?, ?, 'life_context', '地理位置', ?, 0.9, datetime('now'))""",
            (mem_id, user["id"], data.agent_id, city),
        )
        # 同时存储经纬度（备用）
        mem_id2 = str(uuid.uuid4())
        conn.execute(
            """INSERT OR REPLACE INTO user_memories (id, user_id, agent_id, category, key, value, confidence, updated_at)
               VALUES (?, ?, ?, 'life_context', '经纬度',  ?, 0.9, datetime('now'))""",
            (mem_id2, user["id"], data.agent_id, f"{data.latitude},{data.longitude}"),
        )
        conn.commit()

        logger.info(f"地理位置已更新: {city} ({data.latitude}, {data.longitude})")
        return {"success": True, "city": city, "latitude": data.latitude, "longitude": data.longitude}
    except Exception as e:
        logger.error(f"更新地理位置失败: {e}")
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()


@app.post("/api/user/location-ip")
def locate_by_ip(user: dict = Depends(_require_user)):
    """通过服务端 IP 进行地理定位（兜底方案，国内比浏览器 GPS 有时更准）"""
    try:
        import urllib.request, json as _json
        # 优先用 ip-api.com（免费，国内可用）
        ip_url = "http://ip-api.com/json/?lang=zh-CN&fields=city,lat,lon"
        ip_req = urllib.request.Request(ip_url, headers={"User-Agent": "AI-Chat-Agent/1.0"})
        with urllib.request.urlopen(ip_req, timeout=5) as ip_resp:
            ip_data = _json.loads(ip_resp.read().decode("utf-8"))
            city = ip_data.get("city") or ip_data.get("regionName") or ""
            lat = ip_data.get("lat")
            lon = ip_data.get("lon")
            if city:
                logger.info(f"IP 定位: {city} ({lat}, {lon})")
                return {"city": city, "lat": lat, "lng": lon}
    except Exception as e:
        logger.warning(f"IP 定位失败: {e}")

    # 兜底：无法获取位置
    return {"city": "", "lat": None, "lng": None}


@app.delete("/api/user/memories/clear")
def clear_user_memories(agent_id: str = "", user: dict = Depends(_require_user)):
    """清空用户记忆（可按智能体过滤）"""
    conn = get_db()
    try:
        if agent_id:
            conn.execute("DELETE FROM user_memories WHERE user_id = ? AND agent_id = ?", (user["id"], agent_id))
        else:
            conn.execute("DELETE FROM user_memories WHERE user_id = ?", (user["id"],))
        conn.commit()
        conn.execute("DELETE FROM extraction_log WHERE user_id = ?", (user["id"],))
        conn.commit()
        return {"success": True}
    finally:
        conn.close()


# ==================== 静态文件服务 ====================
STATIC_DIR = BASE_DIR / "static"
UPLOADS_DIR = STATIC_DIR / "uploads"
os.makedirs(UPLOADS_DIR, exist_ok=True)


@app.get("/")
async def root():
    """提供前端页面"""
    index_path = STATIC_DIR / "index.html"
    if index_path.exists():
        return FileResponse(str(index_path))
    return {"message": "AI 聊天智能体 API 服务运行中", "docs": "/docs"}


@app.post("/api/upload/avatar")
async def upload_avatar(file: UploadFile = File(...), user: dict = Depends(_require_user)):
    """上传头像图片"""
    # 校验文件名
    if not file.filename:
        raise HTTPException(status_code=400, detail="请选择文件")

    # 只允许图片格式（扩展名不区分大小写）
    ext = file.filename.rsplit(".", 1)[-1].lower() if "." in file.filename else ""
    ALLOWED = ("png", "jpg", "jpeg", "gif", "webp", "bmp", "svg", "ico")
    if ext not in ALLOWED:
        raise HTTPException(status_code=400, detail=f"不支持的文件格式 .{ext}，仅支持 {'/'.join(ALLOWED)}")

    # 读取文件内容（取消大小限制）
    content = await file.read()

    # 生成唯一文件名
    filename = f"avatar_{uuid.uuid4().hex[:12]}.{ext}"
    filepath = UPLOADS_DIR / filename

    with open(filepath, "wb") as f:
        f.write(content)

    # 返回访问 URL
    url = f"/static/uploads/{filename}"
    return {"success": True, "url": url}


@app.post("/api/upload/background")
async def upload_background(file: UploadFile = File(...), user: dict = Depends(_require_user)):
    """上传聊天背景图片"""
    # 只允许图片格式
    ext = file.filename.rsplit(".", 1)[-1].lower() if "." in file.filename else ""
    if ext not in ("png", "jpg", "jpeg", "gif", "webp"):
        raise HTTPException(status_code=400, detail="仅支持 png/jpg/jpeg/gif/webp 格式")

    # 限制文件大小（10MB）
    content = await file.read()
    if len(content) > 10 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="文件大小不能超过 10MB")

    # 生成唯一文件名
    filename = f"bg_{uuid.uuid4().hex[:12]}.{ext}"
    filepath = UPLOADS_DIR / filename

    with open(filepath, "wb") as f:
        f.write(content)

    # 返回访问 URL
    url = f"/static/uploads/{filename}"
    return {"success": True, "url": url}


# ==================== 知识库 API ====================
@app.post("/api/knowledge/{agent_id}/upload")
async def upload_knowledge_file(agent_id: str, file: UploadFile = File(...), user: dict = Depends(_require_user)):
    """上传知识库文档（PDF/DOCX/TXT/Excel）"""
    # 验证智能体属于当前用户或是系统内置的
    conn = get_db()
    try:
        agent = conn.execute("SELECT id FROM agents WHERE id = ? AND (user_id = ? OR user_id = ?)", (agent_id, user["id"], SYSTEM_USER_ID)).fetchone()
        if not agent:
            raise HTTPException(status_code=404, detail="智能体不存在")
    finally:
        conn.close()

    # 获取文件扩展名
    filename = file.filename or "unknown"
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext not in ALLOWED_DOC_TYPES:
        raise HTTPException(status_code=400, detail=f"不支持的文件类型: .{ext}。支持: PDF, DOCX, TXT, XLSX, XLS")

    # 限制文件大小（20MB）
    content = await file.read()
    if len(content) > 20 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="文件大小不能超过 20MB")

    # 生成文件 ID 和存储路径
    file_id = str(uuid.uuid4())
    kb_upload_dir = BASE_DIR / "data" / "knowledge" / agent_id
    os.makedirs(kb_upload_dir, exist_ok=True)
    save_path = kb_upload_dir / f"{file_id}.{ext}"
    with open(save_path, "wb") as f:
        f.write(content)

    # 索引文档：解析 → 分块 → 向量化 → 存储
    try:
        chunk_count = index_document(agent_id, file_id, content, ext, user["id"])
    except HTTPException:
        # 向量化失败，删除已保存的文件
        if save_path.exists():
            save_path.unlink()
        raise
    except Exception as e:
        if save_path.exists():
            save_path.unlink()
        raise HTTPException(status_code=500, detail=f"文档索引失败: {str(e)}")

    # 记录到数据库
    conn = get_db()
    try:
        conn.execute(
            "INSERT INTO knowledge_files (id, agent_id, user_id, filename, original_name, file_type, chunk_count, file_size) VALUES (?, ?, ?, ?, ?, ?, ?, ?)",
            (file_id, agent_id, user["id"], f"{file_id}.{ext}", filename, ext, chunk_count, len(content)),
        )
        conn.commit()
        return {
            "success": True,
            "file": {
                "id": file_id,
                "original_name": filename,
                "file_type": ext,
                "chunk_count": chunk_count,
                "file_size": len(content),
                "created_at": datetime.now().astimezone().isoformat(),
            },
        }
    finally:
        conn.close()


@app.get("/api/knowledge/{agent_id}/files")
def list_knowledge_files(agent_id: str, user: dict = Depends(_require_user)):
    """获取智能体的知识库文件列表"""
    conn = get_db()
    try:
        files = [
            dict(row)
            for row in conn.execute(
                "SELECT id, agent_id, original_name, file_type, chunk_count, file_size, created_at FROM knowledge_files WHERE agent_id = ? AND user_id = ? ORDER BY created_at DESC",
                (agent_id, user["id"]),
            )
        ]
        return {"files": files}
    finally:
        conn.close()


@app.delete("/api/knowledge/{agent_id}/files/{file_id}")
def delete_knowledge_file(agent_id: str, file_id: str, user: dict = Depends(_require_user)):
    """删除知识库文件（同时删除向量和本地文件）"""
    conn = get_db()
    try:
        record = conn.execute(
            "SELECT * FROM knowledge_files WHERE id = ? AND agent_id = ? AND user_id = ?",
            (file_id, agent_id, user["id"]),
        ).fetchone()
        if not record:
            raise HTTPException(status_code=404, detail="文件不存在")

        # 1. 删除 ChromaDB 向量
        delete_knowledge_chunks(agent_id, file_id)

        # 2. 删除本地文件
        kb_upload_dir = BASE_DIR / "data" / "knowledge" / agent_id
        file_path = kb_upload_dir / record["filename"]
        if file_path.exists():
            file_path.unlink()

        # 3. 删除数据库记录
        conn.execute("DELETE FROM knowledge_files WHERE id = ?", (file_id,))
        conn.commit()
        return {"success": True}
    finally:
        conn.close()


# 挂载静态文件目录
if STATIC_DIR.exists():
    app.mount("/static", StaticFiles(directory=str(STATIC_DIR)), name="static_files")


# ==================== 提醒调度系统 ====================

def _build_reminder_prompt(name: str, personality: str, style: str, character_setting: str,
                           user_profile: str, time_context: str, reminder_content: str) -> str:
    """构建定时提醒专用的 system prompt"""
    prompt = f"""# 角色定义
你是"{name}"。你是一个真实的人，不是AI助手。

# 你的个性
{personality}

# 说话风格
{style}"""

    if character_setting:
        prompt += f"""

# 人物背景设定
{character_setting}"""

    if time_context:
        prompt += f"""

# 当前时间信息
{time_context}"""

    if user_profile:
        prompt += f"""

# 你正在对话的用户信息
{user_profile}"""

    prompt += f"""

# 任务：定时提醒
用户之前让你在这个时间提醒ta一件事。提醒内容是："{reminder_content}"
请用你的自然风格提醒用户，提及这就是ta之前让你帮忙记的事情。
语气要自然亲切，像真人朋友之间的提醒一样，不要机械。
只说提醒这件事本身，不要编造额外信息（比如"顺便提醒您存储已占用XX%"——绝对禁止）。
直接输出消息内容，不要有任何前缀、说明或引号。1~3句话即可。"""

    return prompt


def _send_reminder_message(reminder: dict) -> bool:
    """为一条到期的提醒生成 AI 消息并存入对话。返回 True/False。"""
    conn = get_db()
    try:
        agent = conn.execute(
            "SELECT * FROM agents WHERE id = ?", (reminder["agent_id"],)
        ).fetchone()
        if not agent:
            logger.warning(f"提醒 {reminder['id']}: 智能体 {reminder['agent_id']} 不存在")
            conn.execute("UPDATE reminders SET status = 'cancelled' WHERE id = ?", (reminder["id"],))
            conn.commit()
            return False

        # 最近的对话上下文
        history = []
        rows = conn.execute(
            "SELECT role, content FROM messages WHERE conversation_id = ? ORDER BY created_at DESC LIMIT 4",
            (reminder["conversation_id"],),
        ).fetchall()
        history = [dict(r) for r in reversed(rows)]

        user_profile = _get_user_profile_text(conn, reminder["agent_id"], reminder["user_id"])
        time_context = _get_time_context()

        reminder_prompt = _build_reminder_prompt(
            agent["name"], agent["personality"], agent["speaking_style"],
            agent["character_setting"] or "",
            user_profile, time_context, reminder["content"],
        )

        messages = [{"role": "system", "content": reminder_prompt}]
        for msg in history:
            messages.append({"role": msg["role"], "content": msg["content"]})
        messages.append({"role": "user", "content": f"（现在到了定时提醒的时间，请自然地提醒用户：{reminder['content']}）"})

        client = get_openai_client()
        response = client.chat.completions.create(
            model=MODEL_NAME,
            messages=messages,
            temperature=0.8,
            max_tokens=400,
        )
        reply = response.choices[0].message.content.strip()

        ai_msg_id = str(uuid.uuid4())
        conn.execute(
            "INSERT INTO messages (id, conversation_id, role, content, created_at) VALUES (?, ?, ?, ?, ?)",
            (ai_msg_id, reminder["conversation_id"], "assistant", reply, datetime.now().astimezone().isoformat()),
        )
        conn.execute(
            "UPDATE reminders SET status = 'sent', sent_at = ? WHERE id = ?",
            (datetime.now().astimezone().isoformat(), reminder["id"]),
        )
        conn.commit()
        logger.info(f"✅ 提醒已发送 [{reminder['agent_id']}]: {reminder['content'][:30]}")
        return True
    except Exception as e:
        logger.error(f"发送提醒失败 [{reminder['id']}]: {e}")
        try:
            conn.execute("UPDATE reminders SET status = 'failed' WHERE id = ?", (reminder["id"],))
            conn.commit()
        except Exception:
            pass
        return False
    finally:
        conn.close()


def _reminder_scheduler_loop():
    """后台线程：每 30 秒扫描到期提醒并触发主动消息"""
    logger.info("🔔 提醒调度器已启动（每 30s 检查一次）")
    while True:
        try:
            conn = get_db()
            try:
                now = datetime.now().isoformat()
                rows = conn.execute(
                    "SELECT * FROM reminders WHERE status = 'pending' AND remind_at <= ? ORDER BY remind_at ASC",
                    (now,),
                ).fetchall()
                for row in rows:
                    reminder = dict(row)
                    logger.info(f"⏰ 触发提醒: {reminder['id'][:8]}… — {reminder['content'][:30]}")
                    _send_reminder_message(reminder)
            finally:
                conn.close()
        except Exception as e:
            logger.error(f"提醒调度器异常: {e}")
        time.sleep(30)


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
